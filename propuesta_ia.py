import os
import re
from io import BytesIO
from datetime import datetime, timedelta, date
from pathlib import Path

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)


def generar_objetivo_general(nombre_servicio: str) -> str:
    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas comerciales empresariales.

Redacta un párrafo formal, profesional y persuasivo para una propuesta de servicios especializada en: “{nombre_servicio}”.

El texto debe iniciar con la frase:

Es un placer presentarle la siguiente propuesta de servicios especializados en “{nombre_servicio.upper()}”, , cuyo objetivo es...

Después debe desarrollar:

1. Objetivo principal del servicio.
2. Cómo fortalece a la empresa cliente.
3. Qué actividades incluye.
4. Beneficios concretos para la empresa.
5. Lenguaje ejecutivo, serio y corporativo.
6. Todo en un solo párrafo continuo.
7. Redacción elegante y de alto valor percibido.
8. Debe sonar como una propuesta profesional dirigida a empresas en México.

No uses viñetas ni subtítulos. Solo entrega el párrafo final listo para pegar en una propuesta formal.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"



def limpiar_texto(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def extraer_entre_parentesis(texto):
    texto = limpiar_texto(texto)
    encontrados = re.findall(r"\((.*?)\)", texto)
    if encontrados:
        return encontrados[-1].strip()
    return texto


def limpiar_objetivo(texto):
    texto = limpiar_texto(texto)
    return re.sub(r"^Objetivo:\s*", "", texto, flags=re.IGNORECASE).strip()


def formatear_fecha(valor):
    if pd.isna(valor) or valor == "":
        return ""

    if isinstance(valor, datetime):
        return valor.strftime("%d/%m/%Y")

    try:
        fecha = pd.to_datetime(valor, errors="coerce", dayfirst=True)
        if pd.notna(fecha):
            return fecha.strftime("%d/%m/%Y")
    except Exception:
        pass

    return str(valor)

def primer_lunes(anio, mes):
    d = date(anio, mes, 1)
    while d.weekday() != 0:
        d += timedelta(days=1)
    return d


def tercer_lunes(anio, mes):
    return primer_lunes(anio, mes) + timedelta(days=14)


def obtener_dias_festivos_mexico(anio):
    festivos = set()

    festivos.add(date(anio, 1, 1))
    festivos.add(date(anio, 5, 1))
    festivos.add(date(anio, 9, 16))
    festivos.add(date(anio, 12, 25))

    festivos.add(primer_lunes(anio, 2))
    festivos.add(tercer_lunes(anio, 3))
    festivos.add(tercer_lunes(anio, 11))

    return festivos


def es_dia_habil_mexico(fecha):
    if fecha.weekday() == 6:
        return False

    return fecha not in obtener_dias_festivos_mexico(fecha.year)


def restar_15_dias_habiles_mexico(fecha_base):
    contador = 0
    fecha_actual = fecha_base

    while contador < 15:
        fecha_actual -= timedelta(days=1)

        if es_dia_habil_mexico(fecha_actual):
            contador += 1

    return fecha_actual


def convertir_fecha_a_date(valor):
    fecha = pd.to_datetime(valor, errors="coerce", dayfirst=True)

    if pd.isna(fecha):
        return None

    return fecha.date()


def obtener_fecha_legal_desde_registros(registros):
    fechas = []

    for r in registros:
        fecha = convertir_fecha_a_date(r["Fecha"])
        if fecha:
            fechas.append(fecha)

    if not fechas:
        return None

    fecha_mas_antigua = min(fechas)

    return restar_15_dias_habiles_mexico(fecha_mas_antigua)


def formatear_fecha_larga_espanol(fecha):
    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre"
    }

    return f"{fecha.day} de {meses[fecha.month]} del {fecha.year}"

def extraer_datos_excel(archivo_excel):
    df = pd.read_excel(archivo_excel, header=None)

    cliente = ""
    programa = ""
    objetivo = ""

    for i in range(len(df)):
        fila = df.iloc[i].astype(str).tolist()

        for j, valor in enumerate(fila):
            valor_limpio = limpiar_texto(valor)

            if valor_limpio.lower() == "cliente:":
                cliente = limpiar_texto(df.iloc[i, j + 1])

            elif valor_limpio.lower() == "nombre del programa:":
                programa = limpiar_texto(df.iloc[i, j + 1])

            elif valor_limpio.lower().startswith("objetivo:"):
                objetivo = limpiar_objetivo(valor_limpio)

    fila_encabezados = None

    for i in range(len(df)):
        fila = [limpiar_texto(x).lower() for x in df.iloc[i].tolist()]
        if "concepto" in fila and "total factura ($)" in fila and "fecha" in fila:
            fila_encabezados = i
            break

    if fila_encabezados is None:
        raise ValueError("No se encontró la fila de encabezados con Concepto, Total Factura ($) y Fecha.")

    encabezados = [limpiar_texto(x) for x in df.iloc[fila_encabezados].tolist()]

    col_concepto = encabezados.index("Concepto")
    col_monto = encabezados.index("Total Factura ($)")
    col_fecha = encabezados.index("Fecha")

    registros = []

    ultimo_concepto = ""

    for i in range(fila_encabezados + 1, len(df)):
        concepto_raw = limpiar_texto(df.iloc[i, col_concepto])
        monto = df.iloc[i, col_monto]
        fecha = df.iloc[i, col_fecha]

        if concepto_raw:
            ultimo_concepto = extraer_entre_parentesis(concepto_raw)

        if pd.notna(monto):
            registros.append({
                "Metodología / Concepto": ultimo_concepto,
                "Monto": float(monto),
                "Fecha": formatear_fecha(fecha)
            })

    return cliente, programa, objetivo, registros


def generar_introduccion(nombre_servicio: str, conceptos_factura: str) -> str:
    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas de servicios empresariales en México.

Voy a proporcionarte dos parámetros:

1. TEMÁTICA DEL SERVICIO: {nombre_servicio}
2. CONCEPTOS DE FACTURA: {conceptos_factura}

Debes generar únicamente el apartado:

INTRODUCCIÓN

INSTRUCCIONES IMPORTANTES:

- Analiza la temática principal y los conceptos de factura para comprender qué tipo de servicio se prestó.
- No copies literalmente los conceptos.
- Interpreta la información y conviértela en redacción ejecutiva profesional.
- Usa lenguaje corporativo, técnico y formal.
- Redacta un solo párrafo.
- No uses listas.
- No uses subtítulos.
- No uses viñetas.

La introducción debe:

- Explicar por qué ese tema es importante para una empresa.
- Introducir la necesidad del servicio.
- Indicar el propósito del proyecto.
- Mencionar beneficios como crecimiento, eficiencia, posicionamiento, control, rentabilidad o mejora operativa.

Entrega solo el párrafo final.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"

def generar_problematica(nombre_servicio: str, conceptos_factura: str) -> str:
    prompt = f"""
Actúa como consultor corporativo experto en diagnóstico empresarial y propuestas de servicios en México.

Voy a proporcionarte dos parámetros:

1. TEMÁTICA DEL SERVICIO: {nombre_servicio}
2. CONCEPTOS DE FACTURA: {conceptos_factura}

Debes generar únicamente el apartado:

PROBLEMÁTICA

INSTRUCCIONES IMPORTANTES:

- Analiza la temática principal y los conceptos de factura.
- No copies literalmente los conceptos.
- Interpreta la información como riesgos o áreas de oportunidad empresariales.
- Usa lenguaje corporativo, formal y técnico.
- Redacta un solo párrafo.
- No uses listas.
- No uses subtítulos.
- No uses viñetas.

La problemática debe:

- Describir deficiencias o necesidades de una empresa sin ese servicio.
- Exponer consecuencias operativas, comerciales, financieras o estratégicas.
- Justificar la necesidad de apoyo especializado.
- Cerrar indicando la necesidad de implementar soluciones relacionadas con la temática.

Entrega solo el párrafo final.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"

def generar_objetivo_general_tabla(
    nombre_servicio: str,
    conceptos_factura: str,
    nombre_cliente: str
) -> str:

    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas de servicios empresariales en México.

Voy a proporcionarte:

1. TEMÁTICA DEL SERVICIO: {nombre_servicio}
2. CONCEPTOS DE FACTURA: {conceptos_factura}
3. NOMBRE DEL CLIENTE: {nombre_cliente}

Debes generar únicamente el OBJETIVO GENERAL.

INSTRUCCIONES IMPORTANTES:

- Analiza la temática principal para identificar el propósito global del servicio.
- Analiza los conceptos de factura para detectar resultados, actividades, mejoras y beneficios.
- No copies literalmente los conceptos.
- Convierte la información en redacción profesional.
- Usa lenguaje técnico, formal y corporativo.
- Debe sonar como documento dirigido a empresas en México.

REGLAS DE REDACCIÓN:

Debe iniciar exactamente con:

Al finalizar el servicio, {nombre_cliente.upper()}, será capaz de...

Después desarrolla:

- Resultado global que obtendrá la empresa.
- Beneficio principal del servicio.
- Mejoras operativas, comerciales, estratégicas, administrativas o financieras.
- Enfoque de crecimiento, eficiencia, control, posicionamiento, rentabilidad o reducción de riesgos.

FORMATO:

- Un solo párrafo amplio.
- Sin viñetas.
- Sin subtítulos.
- Solo entrega el texto final.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"

def generar_objetivos_especificos(
    nombre_servicio: str,
    conceptos_factura: str,
    nombre_cliente: str
) -> str:

    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas de servicios empresariales en México.

Voy a proporcionarte:

1. TEMÁTICA DEL SERVICIO: {nombre_servicio}
2. CONCEPTOS DE FACTURA: {conceptos_factura}
3. NOMBRE DEL CLIENTE: {nombre_cliente}

Debes generar únicamente los OBJETIVOS ESPECÍFICOS.

INSTRUCCIONES IMPORTANTES:

- Analiza la temática principal y los conceptos de factura.
- No copies literalmente los conceptos.
- Interpreta actividades, beneficios y resultados.
- Genera entre 4 y 6 objetivos específicos.
- Redacción formal, técnica y empresarial.

REGLAS:

Cada objetivo debe iniciar con:

{nombre_cliente.upper()} al término del servicio,...

Deben expresar logros concretos como:

- Diagnosticar
- Diseñar
- Implementar
- Optimizar
- Medir
- Fortalecer
- Incrementar
- Mejorar procesos

FORMATO OBLIGATORIO:

1. Texto...
2. Texto...
3. Texto...
4. Texto...

No agregues títulos.
No agregues explicación adicional.
Solo entrega el resultado final.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"

def generar_metodologia(
    nombre_servicio: str,
    conceptos_factura: str,
) -> str:

    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas de servicios empresariales en México.

Voy a proporcionarte dos parámetros:

1. TEMÁTICA DEL SERVICIO: {nombre_servicio}
2. CONCEPTOS DE FACTURA: {conceptos_factura}

Con esa información debes generar únicamente el apartado:

METODOLOGÍA

INSTRUCCIONES IMPORTANTES:

- La metodología se redacta en función de los conceptos de factura.
- Cada concepto de factura debe convertirse en un punto metodológico desarrollado.
- Conserva la idea principal de cada concepto, pero mejora la redacción con lenguaje profesional, técnico y corporativo.
- No copies solo el concepto; amplíalo explicando qué se analiza, desarrolla, implementa o mejora.
- El texto debe sonar como propuesta formal dirigida a empresas en México.
- No omitas conceptos salvo que sean repetitivos; si lo consideras necesario puedes fusionarlos estratégicamente.

FORMATO OBLIGATORIO:

Primero escribe exactamente:

A continuación, se presentan los temas incluidos en el servicio:

Después desarrolla cada punto en párrafos separados con esta estructura:

**[SUBTÍTULO EN MAYÚSCULAS].** explicación profesional del tema.

REGLAS ESTRICTAS:

- TODOS los subtítulos deben ir completamente en MAYÚSCULAS.
- NO uses viñetas.
- NO uses numeración.
- Cada tema debe ir separado por un salto de línea.
- Mantén redacción clara, elegante y ejecutiva.
- Respeta tildes aun estando en mayúsculas.
- El subtítulo debe terminar en punto dentro de negritas.

EJEMPLO:

**PLAN INTEGRAL PARA GESTIONAR Y MONITOREAR LAS ACTIVIDADES DEL PLAN DE MANTENIMIENTO.** Desarrollo de un esquema operativo que permita organizar, supervisar y dar seguimiento continuo a las acciones preventivas, correctivas y de conservación, asegurando el cumplimiento de tiempos, responsables y objetivos establecidos.

**IDENTIFICACIÓN DE RIESGOS CRÍTICOS QUE PUEDAN AFECTAR LA OPERACIÓN DEL INMUEBLE.** Análisis de condiciones estructurales, técnicas, operativas y ambientales que representen amenazas para la funcionalidad de las instalaciones, con el propósito de prevenir fallas, incidentes y afectaciones a la continuidad operativa.

ENTREGA SOLO EL RESULTADO FINAL.
"""

    try:
        respuesta = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt
        )
        return respuesta.output_text.strip()

    except Exception as e:
        return f"Error: {str(e)}"


def agregar_parrafo(
    documento,
    texto,
    size=12,
    bold=False,
    align=None,
    espacio_despues=12,
    fuente="Times New Roman",
    color=(0, 0, 0),
    line_spacing=1.5,
    variables_negritas=None
):
    import re

    parrafo = documento.add_paragraph()

    if align:
        parrafo.alignment = align

    parrafo.paragraph_format.space_after = Pt(espacio_despues)
    parrafo.paragraph_format.line_spacing = line_spacing

    if variables_negritas is None:
        variables_negritas = []

    # Si no hay variables especiales, comportamiento normal
    if not variables_negritas:
        run = parrafo.add_run(texto)
        run.font.name = fuente
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        return parrafo

    patrones = sorted(
        [re.escape(v) for v in variables_negritas if v],
        key=len,
        reverse=True
    )

    patron = "(" + "|".join(patrones) + ")"
    partes = re.split(patron, texto)

    for parte in partes:
        if not parte:
            continue

        run = parrafo.add_run(parte)
        run.font.name = fuente
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])

        if parte in variables_negritas:
            run.bold = True
        else:
            run.bold = bold

    return parrafo

#CATALOGOS
EMPRESAS = {
    "WAVELENS S.A. DE C.V.": {
        "slug": "wavelens",
        "nombre": "WAVELENS S.A. DE C.V."
    },
    "EMPRESA DEMO S.A. DE C.V.": {
        "slug": "empresa_demo",
        "nombre": "EMPRESA DEMO S.A. DE C.V."
    }
}

CLIENTES = [
    "YUCASA S.A. DE C.V.",
    "CLIENTE DEMO S.A. DE C.V."
]

TIPOGRAFIAS = [
    "Times New Roman",
    "Georgia",
    "Garamond",
    "Cambria",
    "Arial"
]

COLORES = {
    "Naranja": (216, 75, 32),
    "Negro": (0, 0, 0),
    "Rosa pastel": (255, 205, 197),
    "Azul marino": (31, 56, 100),
    "Vino": (112, 48, 48)
}

def crear_word_propuesta(
    plantilla_word,
    nombre_empresa: str,
    nombre_cliente: str,
    lugar_fecha: str,
    nombre_servicio: str,
    texto_propuesta: str,
    introduccion: str,
    problematica: str,
    objetivo_general: str,
    objetivos_especificos: str,
    metodologia: str,
    fuente_titulos: str,
    fuente_texto: str,
    color_titulos: tuple,
    color_texto: tuple,
    color_tabla_encabezado: tuple,
    color_tabla_subtitulos: tuple
) -> BytesIO:

    documento = Document(plantilla_word)

    # Limpia páginas vacías/contenido de la plantilla,
    # pero conserva encabezados, pies y membretes.
    limpiar_contenido_documento(documento)

    # ======================
    # PÁGINA 1 - PORTADA
    # ======================

    for _ in range(5):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        nombre_cliente.upper(),
        size=28,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        espacio_despues=0,
        fuente=fuente_titulos,
        color=color_titulos,
        line_spacing=1.0
    )

    for _ in range(4):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        "PROPUESTA DE",
        size=24,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        espacio_despues=8,
        fuente=fuente_titulos,
        color=(0, 0, 0),
        line_spacing=1.0
    )

    agregar_parrafo(
        documento,
        "SERVICIOS",
        size=24,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        espacio_despues=0,
        fuente=fuente_titulos,
        color=(0, 0, 0),
        line_spacing=1.0
    )

    for _ in range(6):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        nombre_empresa.upper(),
        size=22,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        espacio_despues=0,
        fuente=fuente_titulos,
        color=color_titulos,
        line_spacing=1.0
    )

    documento.add_page_break()

    # ======================
    # PÁGINA 2 - PROPUESTA
    # ======================

    for _ in range(3):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        lugar_fecha,
        size=12,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        fuente=fuente_texto,
        color=color_texto
    )

    for _ in range(2):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        f"Empresa: {nombre_cliente}",
        size=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        fuente=fuente_texto,
        color=color_texto
    )

    for _ in range(2):
        documento.add_paragraph()

    agregar_parrafo(
        documento,
        texto_propuesta,
        size=12,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        fuente=fuente_texto,
        color=color_texto,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente,
            nombre_servicio.upper(),
            nombre_servicio
        ]
    )

    documento.add_page_break()

    # ======================
    # PÁGINA 3 - TABLA DE PROPUESTA
    # ======================

    crear_tabla_propuesta(
        documento=documento,
        nombre_empresa=nombre_empresa,
        nombre_cliente=nombre_cliente,
        nombre_servicio=nombre_servicio,
        introduccion=introduccion,
        problematica=problematica,
        objetivo_general=objetivo_general,
        objetivos_especificos=objetivos_especificos,
        metodologia=metodologia,
        fuente_texto=fuente_texto,
        color_encabezado=color_tabla_encabezado,
        color_subtitulos=color_tabla_subtitulos
    )

    # ======================
    # PÁGINA - AVISO DE CONFIDENCIALIDAD
    # ======================

    documento.add_page_break()

    texto_confidencialidad = generar_aviso_confidencialidad(
        nombre_empresa,
        nombre_cliente
    )

    agregar_parrafo(
        documento,
        "AVISO DE CONFIDENCIALIDAD Y PRIVACIDAD",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        fuente=fuente_titulos,
        color=color_titulos,
        espacio_despues=18
    )

    agregar_parrafo(
        documento,
        texto_confidencialidad,
        size=12,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        fuente=fuente_texto,
        color=color_texto,
        line_spacing=1.5,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    documento.add_paragraph()

    agregar_parrafo(
        documento,
        "ACEPTADO:",
        size=12,
        fuente=fuente_texto
    )

    documento.add_paragraph()

    agregar_parrafo(
        documento,
        "FIRMA: ____________________",
        size=12,
        fuente=fuente_texto
    )

    documento.add_paragraph()

    agregar_parrafo(
        documento,
        "PUESTO: ___________________",
        size=12,
        fuente=fuente_texto
    )

    archivo = BytesIO()
    documento.save(archivo)
    archivo.seek(0)

    return archivo

def mostrar_modulo_propuesta_ia():
    st.title("Generador de propuesta con IA")

    empresa_seleccionada = st.selectbox(
        "Selecciona la empresa que presta el servicio",
        list(EMPRESAS.keys())
    )

    nombre_empresa = EMPRESAS[empresa_seleccionada]["nombre"]
    empresa_slug = EMPRESAS[empresa_seleccionada]["slug"]

    nombre_cliente = st.selectbox(
        "Selecciona el cliente",
        CLIENTES
    )

    nombre_servicio = st.text_input(
        "Nombre del servicio",
        placeholder="Ej. CONSULTORÍA EN ESTRATEGIAS DE COMERCIALIZACIÓN"
    )

    archivo_excel = st.file_uploader(
        "Sube Excel para obtener conceptos de factura",
        type=["xlsx", "xls"]
    )

    conceptos_factura = ""

    if archivo_excel is not None:
        try:
            cliente_excel, programa_excel, objetivo_excel, registros = extraer_datos_excel(archivo_excel)
            fecha_legal = obtener_fecha_legal_desde_registros(registros)

            if fecha_legal:
                lugar_fecha = f"Ciudad de Mérida, Yucatán a {formatear_fecha_larga_espanol(fecha_legal)}"
            else:
                lugar_fecha = "Ciudad de Mérida, Yucatán"

            conceptos_unicos = []
            vistos = set()

            for r in registros:
                concepto = r["Metodología / Concepto"].strip()

                if concepto and concepto not in vistos:
                    vistos.add(concepto)
                    conceptos_unicos.append(concepto)

            conceptos_factura = "\n".join(
                [f"{i + 1}. {c}" for i, c in enumerate(conceptos_unicos)]
            )

            st.subheader("Conceptos detectados desde Excel")
            st.text_area(
                "Vista previa",
                conceptos_factura,
                height=180
            )
            st.write("**Fecha automática detectada:**")
            st.write(lugar_fecha)

        except Exception as e:
            st.error(f"Error al leer Excel: {e}")

    fuente_titulos = st.selectbox(
        "Tipografía para títulos",
        TIPOGRAFIAS,
        index=1
    )

    fuente_texto = st.selectbox(
        "Tipografía para texto",
        TIPOGRAFIAS,
        index=0
    )

    color_titulo_nombre = st.selectbox(
        "Color para nombres en portada",
        list(COLORES.keys()),
        index=0
    )

    color_texto_documento = st.selectbox(
        "Color del texto del documento",
        list(COLORES.keys()),
        index=1
    )

    color_titulos = COLORES[color_titulo_nombre]
    color_texto = COLORES[color_texto_documento]

    color_encabezado_tabla_nombre = st.selectbox(
        "Color de encabezados de tabla",
        list(COLORES.keys()),
        index=0
    )

    color_subtitulos_tabla_nombre = st.selectbox(
        "Color de subtítulos de tabla",
        list(COLORES.keys()),
        index=2
    )

    color_tabla_encabezado = COLORES[color_encabezado_tabla_nombre]
    color_tabla_subtitulos = COLORES[color_subtitulos_tabla_nombre]

    if st.button("Generar propuesta"):

        if not nombre_servicio.strip():
            st.warning("Escribe el nombre del servicio.")
            return

        if not conceptos_factura.strip():
            st.warning("Sube un Excel con conceptos de factura.")
            return

        ruta_plantilla = Path("membretes") / empresa_slug / f"plantilla_{empresa_slug}.docx"

        if not ruta_plantilla.exists():
            st.error("No existe la plantilla Word de esta empresa.")
            return

        with st.spinner("Generando textos con IA..."):
            texto_propuesta = generar_objetivo_general(nombre_servicio)
            introduccion = generar_introduccion(nombre_servicio, conceptos_factura)
            problematica = generar_problematica(nombre_servicio, conceptos_factura)
            objetivo_general = generar_objetivo_general_tabla(nombre_servicio, conceptos_factura, nombre_cliente)
            objetivos_especificos = generar_objetivos_especificos(nombre_servicio, conceptos_factura, nombre_cliente)
            metodologia = generar_metodologia(nombre_servicio, conceptos_factura)
        st.subheader("Texto generado para la página 2")
        st.write(texto_propuesta)

        st.subheader("Introducción generada")
        st.write(introduccion)

        st.subheader("Problemática generada")
        st.write(problematica)

        archivo_word = crear_word_propuesta(
            plantilla_word=str(ruta_plantilla),
            nombre_empresa=nombre_empresa,
            nombre_cliente=nombre_cliente,
            lugar_fecha=lugar_fecha,
            nombre_servicio=nombre_servicio,
            texto_propuesta=texto_propuesta,
            introduccion=introduccion,
            problematica=problematica,
            objetivo_general=objetivo_general,
            objetivos_especificos=objetivos_especificos,
            metodologia=metodologia,
            fuente_titulos=fuente_titulos,
            fuente_texto=fuente_texto,
            color_titulos=color_titulos,
            color_texto=color_texto,
            color_tabla_encabezado=color_tabla_encabezado,
            color_tabla_subtitulos=color_tabla_subtitulos
        )

        st.download_button(
            label="Descargar Word",
            data=archivo_word,
            file_name="propuesta_servicio.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def limpiar_contenido_documento(documento):
    body = documento._element.body

    for elemento in list(body):
        if elemento.tag.endswith("sectPr"):
            continue
        body.remove(elemento)


def rgb_a_hex(color):
    return "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])


def colorear_celda(celda, color):
    tc_pr = celda._tc.get_or_add_tcPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:fill"), rgb_a_hex(color))
    tc_pr.append(sombreado)


def configurar_texto_celda(
    celda,
    texto,
    fuente="Times New Roman",
    size=11,
    bold=False,
    color=(0, 0, 0),
    align=WD_ALIGN_PARAGRAPH.CENTER,
    variables_negritas=None
):
    import re

    celda.text = ""

    parrafo = celda.paragraphs[0]
    parrafo.alignment = align

    if variables_negritas is None:
        variables_negritas = []

    if not variables_negritas:
        run = parrafo.add_run(texto)
        run.font.name = fuente
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
    else:
        patrones = sorted(
            [re.escape(v) for v in variables_negritas if v],
            key=len,
            reverse=True
        )

        patron = "(" + "|".join(patrones) + ")"
        partes = re.split(patron, texto)

        for parte in partes:
            if not parte:
                continue

            run = parrafo.add_run(parte)
            run.font.name = fuente
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])

            if parte in variables_negritas:
                run.bold = True
            else:
                run.bold = bold

    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def crear_tabla_propuesta(
    documento,
    nombre_empresa,
    nombre_cliente,
    nombre_servicio,
    introduccion,
    problematica,
    objetivo_general,
    objetivos_especificos,
    metodologia,
    fuente_texto,
    color_encabezado,
    color_subtitulos
):
    tabla = documento.add_table(rows=10, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    poner_bordes_tabla(tabla)

    # Anchos aproximados
    for row in tabla.rows:
        row.cells[0].width = Inches(2.8)
        row.cells[1].width = Inches(4.4)

    # Fila 1: encabezado principal
    celda_titulo = tabla.cell(0, 0).merge(tabla.cell(0, 1))
    colorear_celda(celda_titulo, color_encabezado)
    configurar_texto_celda(
        celda_titulo,
        "PROPUESTA DE SERVICIO",
        fuente=fuente_texto,
        size=11,
        bold=True,
        color=(0, 0, 0)
    )

    # Fila 2: Nombre empresa
    colorear_celda(tabla.cell(1, 0), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(1, 0),
        "NOMBRE DE EMPRESA:",
        fuente=fuente_texto,
        bold=True
    )
    configurar_texto_celda(
        tabla.cell(1, 1),
        nombre_empresa.upper(),
        fuente=fuente_texto
    )

    # Fila 3: Nombre cliente
    colorear_celda(tabla.cell(2, 0), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(2, 0),
        "NOMBRE DE CLIENTE:",
        fuente=fuente_texto,
        bold=True
    )
    configurar_texto_celda(
        tabla.cell(2, 1),
        nombre_cliente.upper(),
        fuente=fuente_texto
    )

    # Fila 4: Servicio
    celda_servicio = tabla.cell(3, 0).merge(tabla.cell(3, 1))
    colorear_celda(celda_servicio, color_encabezado)
    configurar_texto_celda(
        celda_servicio,
        nombre_servicio.upper(),
        fuente=fuente_texto,
        bold=True
    )

    # Fila 5: Introducción
    colorear_celda(tabla.cell(4, 0), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(4, 0),
        "INTRODUCCIÓN:",
        fuente=fuente_texto,
        bold=True
    )
    configurar_texto_celda(
        tabla.cell(4, 1),
        introduccion,
        fuente=fuente_texto,
        size=10,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    # Fila 6: Problemática
    colorear_celda(tabla.cell(5, 0), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(5, 0),
        "PROBLEMÁTICA:",
        fuente=fuente_texto,
        bold=True
    )
    configurar_texto_celda(
        tabla.cell(5, 1),
        problematica,
        fuente=fuente_texto,
        size=10,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    # Fila 7: Encabezado OBJETIVOS
    celda_objetivos = tabla.cell(6, 0).merge(tabla.cell(6, 1))
    colorear_celda(celda_objetivos, color_encabezado)
    configurar_texto_celda(
        celda_objetivos,
        "OBJETIVOS",
        fuente=fuente_texto,
        size=11,
        bold=True,
        color=(0, 0, 0)
    )

    # Fila 8: GENERAL / ESPECÍFICOS
    colorear_celda(tabla.cell(7, 0), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(7, 0),
        "GENERAL",
        fuente=fuente_texto,
        size=11,
        bold=True
    )

    colorear_celda(tabla.cell(7, 1), color_subtitulos)
    configurar_texto_celda(
        tabla.cell(7, 1),
        "ESPECÍFICOS",
        fuente=fuente_texto,
        size=11,
        bold=True
    )

    # Fila 9: Objetivo general / Objetivos específicos
    configurar_texto_celda(
        tabla.cell(8, 0),
        objetivo_general,
        fuente=fuente_texto,
        size=10,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    configurar_texto_celda(
        tabla.cell(8, 1),
        objetivos_especificos,
        fuente=fuente_texto,
        size=10,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    # Fila 10: Metodología
    colorear_celda(tabla.cell(9, 0), color_subtitulos)

    configurar_texto_celda(
        tabla.cell(9, 0),
        "METODOLOGÍA",
        fuente=fuente_texto,
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER
    )

    configurar_texto_celda(
        tabla.cell(9, 1),
        metodologia,
        fuente=fuente_texto,
        size=10,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        variables_negritas=[
            nombre_empresa,
            nombre_cliente
        ]
    )

    return tabla

def poner_bordes_tabla(tabla):
    tbl = tabla._tbl
    tbl_pr = tbl.tblPr

    bordes = OxmlElement("w:tblBorders")

    for borde in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elemento = OxmlElement(f"w:{borde}")
        elemento.set(qn("w:val"), "single")
        elemento.set(qn("w:sz"), "8")
        elemento.set(qn("w:space"), "0")
        elemento.set(qn("w:color"), "000000")
        bordes.append(elemento)

    tbl_pr.append(bordes)


def generar_aviso_confidencialidad(nombre_empresa, nombre_cliente):
    texto = f"""
{nombre_empresa}, no divulgará a terceros la información confidencial perteneciente a {nombre_cliente} ni utilizará dicha información para otro fin que no esté relacionado con la prestación de los servicios, según se acuerda esta Carta Convenio, o conforme a las autorizaciones posteriores.

Cuando la relación entre las partes haya concluido, {nombre_empresa}, destruirá o devolverá a la compañía, según se indique por escrito, cualquier información confidencial en formato impreso o electrónico que haya sido recibido como consecuencia de la prestación de los servicios, excepto aquellos documentos electrónicos que deban ser conservados en cumplimiento de las disposiciones legal o profesionales vigentes.

Para la interpretación, ejecución y cumplimiento de esta carta convenio, las partes se someten a los tribunales de la Ciudad de Mérida, Yucatán y a las leyes competentes, renunciando expresamente a cualquier otro fuero que alguna de las partes tenga o llegase a tener por razón de su domicilio presente o futuro o por cualquier otra razón. 

La presente constituye la totalidad del acuerdo y el entendimiento entre, {nombre_empresa}, y {nombre_cliente}, con respecto al tema que nos concierne, esta carta anula todos los acuerdos y entendimientos anteriores que se hubieren llevado a cabo entre las partes relativas con el objeto de la presente, la cual dejará de tener cualquier fuerza o efecto. Cualquier modificación de los términos de la presente, se hará por escrito y no será válida a menos que se haya firmado por un representante debidamente autorizado de la parte.

Acepto las condiciones de esta carta como un acuerdo de voluntades entre {nombre_empresa}, por una parte, {nombre_cliente}, por la otra, cuyas condiciones y disposiciones he leído y entendido plenamente.
"""
    return texto.strip()