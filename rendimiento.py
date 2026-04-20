import random
from io import BytesIO

import streamlit as st
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
import os

THIN_BLACK = Side(style="thin", color="000000")


def normalizar_texto(valor):
    if valor is None:
        return ""
    return str(valor).strip()


def centrar(cell, bold=False, size=12):
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.font = Font(bold=bold, size=size)


def aplicar_borde_rango(ws, r1, c1, r2, c2):
    for row in ws.iter_rows(min_row=r1, max_row=r2, min_col=c1, max_col=c2):
        for cell in row:
            cell.border = Border(
                left=THIN_BLACK,
                right=THIN_BLACK,
                top=THIN_BLACK,
                bottom=THIN_BLACK
            )


def ajustar_ancho(ws, col_idx, width):
    ws.column_dimensions[get_column_letter(col_idx)].width = width


def limpiar_hoja_si_existe(wb, nombre_hoja):
    if nombre_hoja in wb.sheetnames:
        del wb[nombre_hoja]


def safe_sheet_name(name):
    invalid = ['\\', '/', '*', '[', ']', ':', '?']
    for ch in invalid:
        name = name.replace(ch, "")
    return name[:31]


def generar_desempeno():
    return random.choices(
        population=[5, 4, 3],
        weights=[60, 35, 5],
        k=1
    )[0]


def primer_no_vacio_en_fila(ws, row_num):
    for cell in ws[row_num]:
        if normalizar_texto(cell.value):
            return normalizar_texto(cell.value)
    return ""


def color_hex_desde_fill(cell):
    fill = cell.fill
    if fill is None:
        return "D9D9D9"

    fg = fill.fgColor
    if fg is None:
        return "D9D9D9"

    rgb = getattr(fg, "rgb", None)
    if rgb:
        rgb = str(rgb)
        if len(rgb) == 8:
            return rgb[2:]
        if len(rgb) == 6:
            return rgb

    return "D9D9D9"


def es_participa(cell):
    valor = cell.value
    s = normalizar_texto(valor).strip().lower()
    return s == "a"


def leer_tabla_actividades(file_bytes):
    wb = load_workbook(BytesIO(file_bytes))
    ws = wb.active

    titulo = primer_no_vacio_en_fila(ws, 1) or "TABLA DE ACTIVIDADES"
    puesto = primer_no_vacio_en_fila(ws, 2) or "PUESTO SIN DEFINIR"

    header_row = 3
    start_data_row = 5
    col_num = 1
    col_act = 2
    start_emp_col = 3

    empleados = []
    col = start_emp_col

    while col <= ws.max_column:
        nombre = normalizar_texto(ws.cell(header_row, col).value)
        if nombre == "":
            break

        color = color_hex_desde_fill(ws.cell(header_row, col))

        empleados.append({
            "nombre": nombre,
            "col": col,
            "color": color
        })
        col += 1

    actividades = []
    row = start_data_row

    while row <= ws.max_row:
        num = ws.cell(row, col_num).value
        actividad = normalizar_texto(ws.cell(row, col_act).value)

        if actividad == "":
            break

        participantes = []
        for emp in empleados:
            cell_part = ws.cell(row, emp["col"])
            if es_participa(cell_part):
                participantes.append({
                    "nombre": emp["nombre"],
                    "color": emp["color"]
                })

        actividades.append({
            "numero": num,
            "actividad": actividad,
            "participantes": participantes
        })

        row += 1

    return {
        "titulo": titulo,
        "puesto": puesto,
        "empleados": empleados,
        "actividades": actividades
    }


def generar_rendimiento(actividades, semanas=4):
    resultado = []

    for act in actividades:
        semanas_data = []
        for semana in range(1, semanas + 1):
            bloques = []
            for participante in act["participantes"]:
                bloques.append({
                    "nombre": participante["nombre"],
                    "color": participante["color"],
                    "valor": generar_desempeno()
                })

            semanas_data.append({
                "semana": semana,
                "bloques": bloques
            })

        resultado.append({
            "numero": act["numero"],
            "actividad": act["actividad"],
            "semanas": semanas_data
        })

    return resultado


def render_preview_html(puesto, mes_anio, rendimiento):
    html = f"""
    <div style="font-family:Arial,sans-serif; overflow-x:auto; margin-bottom:30px;">
      <table style="border-collapse:collapse; min-width:1100px; width:100%;">
        <tr>
          <th colspan="5" style="border:1px solid #000; background:#f2be00; font-size:22px; padding:8px;">
            TABLA DE RENDIMIENTO
          </th>
        </tr>
        <tr>
          <th colspan="5" style="border:1px solid #000; background:#f2be00; font-size:20px; padding:8px;">
            {puesto}
          </th>
        </tr>
        <tr>
          <th colspan="5" style="border:1px solid #000; background:#f2be00; font-size:20px; padding:8px;">
            {mes_anio}
          </th>
        </tr>
        <tr>
          <th style="border:1px solid #000; background:#eeeeee; width:100px; font-size:18px; padding:6px;">
            #ACT.
          </th>
          <th style="border:1px solid #000; background:#eeeeee; width:240px; font-size:18px;">
            SEMANA 1
          </th>
          <th style="border:1px solid #000; background:#eeeeee; width:240px; font-size:18px;">
            SEMANA 2
          </th>
          <th style="border:1px solid #000; background:#eeeeee; width:240px; font-size:18px;">
            SEMANA 3
          </th>
          <th style="border:1px solid #000; background:#eeeeee; width:240px; font-size:18px;">
            SEMANA 4
          </th>
        </tr>
    """

    for item in rendimiento:
        html += f"""
        <tr>
          <td style="border:1px solid #000; background:#FFFFFF; text-align:center; font-size:18px; padding:6px;">
            {item['numero']}
          </td>
        """

        for semana in item["semanas"]:
            bloques = semana["bloques"]

            if not bloques:
                html += '<td style="border:1px solid #000; height:38px; background:#fff;"></td>'
                continue

            html += '<td style="border:1px solid #000; padding:0; height:38px;">'
            html += '<div style="display:flex; width:100%; height:38px;">'

            for i, bloque in enumerate(bloques):
                borde = "border-right:1px solid #000;" if i < len(bloques) - 1 else ""
                html += f"""
                <div style="
                    flex:1;
                    background:#{bloque['color']};
                    {borde}
                    display:flex;
                    justify-content:center;
                    align-items:center;
                    font-size:18px;
                ">
                    {bloque['valor']}
                </div>
                """

            html += '</div></td>'

        html += "</tr>"

    html += "</table></div>"
    return html


def escribir_hoja_rendimiento(wb, nombre_hoja, puesto, mes_anio, rendimiento):
    limpiar_hoja_si_existe(wb, nombre_hoja)
    ws = wb.create_sheet(title=safe_sheet_name(nombre_hoja))

    yellow_fill = PatternFill(fill_type="solid", fgColor="F2BE00")
    gray_fill = PatternFill(fill_type="solid", fgColor="EFEFEF")
    white_fill = PatternFill(fill_type="solid", fgColor="FFFFFF")

    bloques_semana = {
        1: (2, 7),
        2: (8, 13),
        3: (14, 19),
        4: (20, 25)
    }

    ws.merge_cells("A1:Y1")
    ws["A1"] = "TABLA DE RENDIMIENTO"
    ws["A1"].fill = yellow_fill
    centrar(ws["A1"], bold=True, size=16)

    ws.merge_cells("A2:Y2")
    ws["A2"] = puesto
    ws["A2"].fill = yellow_fill
    centrar(ws["A2"], bold=True, size=15)

    ws.merge_cells("A3:Y3")
    ws["A3"] = mes_anio
    ws["A3"].fill = yellow_fill
    centrar(ws["A3"], bold=True, size=15)

    ws.merge_cells("A4:A5")
    ws["A4"] = "#ACT."
    ws["A4"].fill = gray_fill
    centrar(ws["A4"], bold=True, size=14)

    for semana, (c1, c2) in bloques_semana.items():
        ws.merge_cells(start_row=4, start_column=c1, end_row=4, end_column=c2)
        cell = ws.cell(4, c1)
        cell.value = f"SEMANA {semana}"
        cell.fill = gray_fill
        centrar(cell, bold=True, size=14)

        for col in range(c1, c2 + 1):
            ws.cell(5, col).fill = white_fill

    aplicar_borde_rango(ws, 1, 1, 5, 25)

    ajustar_ancho(ws, 1, 10)
    for col in range(2, 26):
        ajustar_ancho(ws, col, 5)

    ws.row_dimensions[1].height = 30
    ws.row_dimensions[2].height = 28
    ws.row_dimensions[3].height = 28
    ws.row_dimensions[4].height = 26
    ws.row_dimensions[5].height = 12

    start_row = 6

    for idx, item in enumerate(rendimiento, start=0):
        excel_row = start_row + idx
        ws.row_dimensions[excel_row].height = 28

        c = ws.cell(excel_row, 1)
        c.value = item["numero"]
        c.fill = white_fill
        centrar(c, size=14)
        c.border = Border(left=THIN_BLACK, right=THIN_BLACK, top=THIN_BLACK, bottom=THIN_BLACK)

        for sem_idx, semana in enumerate(item["semanas"], start=1):
            c1, c2 = bloques_semana[sem_idx]
            bloques = semana["bloques"]

            if len(bloques) == 0:
                ws.merge_cells(start_row=excel_row, start_column=c1, end_row=excel_row, end_column=c2)
                mc = ws.cell(excel_row, c1)
                mc.value = ""
                mc.fill = white_fill
                centrar(mc, size=14)
                aplicar_borde_rango(ws, excel_row, c1, excel_row, c2)

            elif len(bloques) == 1:
                b = bloques[0]
                ws.merge_cells(start_row=excel_row, start_column=c1, end_row=excel_row, end_column=c2)
                mc = ws.cell(excel_row, c1)
                mc.value = b["valor"]
                mc.fill = PatternFill(fill_type="solid", fgColor=b["color"])
                centrar(mc, size=14)
                aplicar_borde_rango(ws, excel_row, c1, excel_row, c2)

            elif len(bloques) == 2:
                rangos = [(c1, c1 + 2), (c1 + 3, c2)]
                for b, (r1, r2) in zip(bloques, rangos):
                    ws.merge_cells(start_row=excel_row, start_column=r1, end_row=excel_row, end_column=r2)
                    mc = ws.cell(excel_row, r1)
                    mc.value = b["valor"]
                    mc.fill = PatternFill(fill_type="solid", fgColor=b["color"])
                    centrar(mc, size=14)
                    aplicar_borde_rango(ws, excel_row, r1, excel_row, r2)

            else:
                total_cols = c2 - c1 + 1
                n = len(bloques)
                ancho_base = total_cols // n
                resto = total_cols % n

                rangos = []
                actual = c1

                for i in range(n):
                    ancho = ancho_base + (1 if i < resto else 0)
                    fin = actual + ancho - 1
                    rangos.append((actual, fin))
                    actual = fin + 1

                for b, (r1, r2) in zip(bloques, rangos):
                    ws.merge_cells(start_row=excel_row, start_column=r1, end_row=excel_row, end_column=r2)
                    mc = ws.cell(excel_row, r1)
                    mc.value = b["valor"]
                    mc.fill = PatternFill(fill_type="solid", fgColor=b["color"])
                    centrar(mc, size=14)
                    aplicar_borde_rango(ws, excel_row, r1, excel_row, r2)

    return wb


def construir_excel_final(file_bytes, puesto, mes1, anio1, mes2, anio2, rendimiento1, rendimiento2):
    wb = load_workbook(BytesIO(file_bytes))

    hoja1 = f"RENDIMIENTO {mes1} {anio1}"
    hoja2 = f"RENDIMIENTO {mes2} {anio2}"

    escribir_hoja_rendimiento(
        wb=wb,
        nombre_hoja=hoja1,
        puesto=puesto,
        mes_anio=f"{mes1} {anio1}",
        rendimiento=rendimiento1
    )

    escribir_hoja_rendimiento(
        wb=wb,
        nombre_hoja=hoja2,
        puesto=puesto,
        mes_anio=f"{mes2} {anio2}",
        rendimiento=rendimiento2
    )

    salida = BytesIO()
    wb.save(salida)
    salida.seek(0)
    return salida.getvalue()


def mostrar_modulo_rendimiento():
    st.title("Generador de dos tablas de rendimiento desde Excel")

    uploaded_file = st.file_uploader(
        "Sube el Excel con la TABLA DE ACTIVIDADES",
        type=["xlsx"],
        key="rendimiento_uploader"
    )

    meses = [
        "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
        "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"
    ]

    st.subheader("Mes 1")
    c1, c2 = st.columns(2)
    with c1:
        mes1 = st.selectbox("Mes 1", meses, index=4, key="mes1")
    with c2:
        anio1 = st.number_input("Año 1", min_value=2024, max_value=2100, value=2025, step=1, key="anio1")

    st.subheader("Mes 2")
    c3, c4 = st.columns(2)
    with c3:
        mes2 = st.selectbox("Mes 2", meses, index=5, key="mes2")
    with c4:
        anio2 = st.number_input("Año 2", min_value=2024, max_value=2100, value=2025, step=1, key="anio2")

    if uploaded_file is not None:
        file_bytes = uploaded_file.read()

        try:
            data = leer_tabla_actividades(file_bytes)

            puesto = data["puesto"]
            empleados = data["empleados"]
            actividades = data["actividades"]

            st.success("Archivo leído correctamente.")
            st.write(f"**Puesto detectado:** {puesto}")
            st.write(f"**Empleados detectados:** {len(empleados)}")
            st.write(f"**Actividades detectadas:** {len(actividades)}")

            with st.expander("Ver empleados y colores detectados"):
                for emp in empleados:
                    st.markdown(
                        f"""
                        <div style="display:flex; align-items:center; gap:10px; margin-bottom:6px;">
                            <div style="width:22px; height:22px; background:#{emp['color']}; border:1px solid #000;"></div>
                            <div>{emp['nombre']}</div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            with st.expander("Ver participación detectada por actividad"):
                for act in actividades:
                    nombres = [p["nombre"] for p in act["participantes"]]
                    st.write(f"{act['numero']}. {act['actividad']} → {', '.join(nombres) if nombres else 'Sin participantes'}")

            col_btn1, col_btn2 = st.columns(2)

            with col_btn1:
                if st.button("Generar dos meses", key="btn_generar"):
                    st.session_state["rendimiento_mes_1"] = generar_rendimiento(actividades, semanas=4)
                    st.session_state["rendimiento_mes_2"] = generar_rendimiento(actividades, semanas=4)
                    st.session_state["file_bytes_original"] = file_bytes
                    st.session_state["puesto_detectado"] = puesto

            with col_btn2:
                if st.button("Regenerar aleatoriamente", key="btn_regenerar"):
                    st.session_state["rendimiento_mes_1"] = generar_rendimiento(actividades, semanas=4)
                    st.session_state["rendimiento_mes_2"] = generar_rendimiento(actividades, semanas=4)

            if "rendimiento_mes_1" in st.session_state and "rendimiento_mes_2" in st.session_state:
                st.subheader("Vista previa: Mes 1")
                html1 = render_preview_html(
                    puesto=st.session_state["puesto_detectado"],
                    mes_anio=f"{mes1} {anio1}",
                    rendimiento=st.session_state["rendimiento_mes_1"]
                )
                st.components.v1.html(html1, height=700, scrolling=True)

                st.subheader("Vista previa: Mes 2")
                html2 = render_preview_html(
                    puesto=st.session_state["puesto_detectado"],
                    mes_anio=f"{mes2} {anio2}",
                    rendimiento=st.session_state["rendimiento_mes_2"]
                )
                st.components.v1.html(html2, height=700, scrolling=True)

                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
                    tmp_doc_path = tmp_doc.name

                st.write("Entrando a generar Word...")

                try:
                    with st.spinner("Generando Word..."):
                        construir_word_final(
                            puesto=st.session_state["puesto_detectado"],
                            mes1=mes1,
                            anio1=anio1,
                            mes2=mes2,
                            anio2=anio2,
                            rendimiento1=st.session_state["rendimiento_mes_1"],
                            rendimiento2=st.session_state["rendimiento_mes_2"],
                            output_file=tmp_doc_path
                        )

                    st.write("Word generado correctamente")

                    with open(tmp_doc_path, "rb") as f:
                        word_bytes = f.read()

                    st.download_button(
                        label="Descargar Word final con dos tablas",
                        data=word_bytes,
                        file_name=f"rendimiento_{mes1.lower()}_{anio1}_{mes2.lower()}_{anio2}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Error al generar el Word: {type(e).__name__}: {e}")

                finally:
                    if os.path.exists(tmp_doc_path):
                        os.remove(tmp_doc_path)

        except Exception as e:
            st.error(f"Ocurrió un error al procesar el archivo: {e}")
    else:
        st.info("Sube el Excel de actividades para comenzar.")



def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def center_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")


def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color_hex)


def set_cell_border(cell, top="000000", bottom="000000", left="000000", right="000000", size="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    edges = {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }

    for edge, color in edges.items():
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def style_word_cell(
    cell,
    text="",
    bg_color="FFFFFF",
    font_color="000000",
    bold=False,
    size=10,
    align="center"
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(font_color)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_background(cell, bg_color)
    set_cell_border(cell)


def configure_document_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)


def add_rendimiento_table_to_doc(doc, puesto, mes_anio, rendimiento):
    # Título fuera de la tabla
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TABLAS DE RENDIMIENTO DEL PUESTO")
    run.bold = True
    run.font.size = Pt(17)
    run.font.name = "Arial"

    doc.add_paragraph("")

    # Estructura:
    # col 0 = # actividad
    # cada semana = 6 columnas
    # y entre semanas una columna separadora blanca (excepto al final)
    total_cols = 1 + 1 + 6 + 1 + 6 + 1 + 6 + 1 + 6  # 28 columnas
    total_rows = 4 + len(rendimiento)

    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    center_table(table)

    # Anchos
    # primera columna más angosta, semanas amplias, separadores pequeños
    col_widths = [2.8, 0.45]  # más ancha actividad

    for i in range(4):
        col_widths += [0.78] * 6
        if i < 3:
            col_widths += [0.45]

    for row in table.rows:
        for i, w in enumerate(col_widths):
            set_cell_width(row.cells[i], w)

    # Alturas
    set_row_height(table.rows[0], 520)
    set_row_height(table.rows[1], 520)
    set_row_height(table.rows[2], 520)
    set_row_height(table.rows[3], 520)

    yellow = "F2BE00"
    gray = "EFEFEF"
    white = "FFFFFF"

    # Fila 0
    c = table.cell(0, 0)
    c.merge(table.cell(0, total_cols - 1))
    style_word_cell(c, "TABLA DE RENDIMIENTO", bg_color=yellow, bold=True, size=15)

    # Fila 1
    c = table.cell(1, 0)
    c.merge(table.cell(1, total_cols - 1))
    style_word_cell(c, puesto, bg_color=yellow, bold=True, size=14)

    # Fila 2
    c = table.cell(2, 0)
    c.merge(table.cell(2, total_cols - 1))
    style_word_cell(c, mes_anio, bg_color=yellow, bold=True, size=14)

    # Encabezado columna actividad
    act_cell = table.cell(3, 0)
    style_word_cell(
        act_cell,
        "#\nACTIVIDAD",
        bg_color=white,
        bold=True,
        size=11
    )    # Semanas
    week_starts = [2, 9, 16, 23]
    separator_cols = [1, 8, 15, 22]

    for idx, start_col in enumerate(week_starts, start=1):
        end_col = start_col + 5
        week_cell = table.cell(3, start_col)
        week_cell.merge(table.cell(3, end_col))
        style_word_cell(
            week_cell,
            f"SEMANA {idx}",
            bg_color=white,
            bold=True,
            size=12
        )

    # Separadores blancos SOLO desde la fila de encabezado de semanas hacia abajo
    for sep_col in separator_cols:
        for r in range(3, total_rows):
            cell = table.cell(r, sep_col)
            cell.text = ""
            set_cell_background(cell, "FFFFFF")
            set_cell_border(cell, top="FFFFFF", bottom="FFFFFF", left="FFFFFF", right="FFFFFF")

    # Filas de datos
    data_start = 4
    for i, item in enumerate(rendimiento):
        row_idx = data_start + i
        set_row_height(table.rows[row_idx], 420)

        style_word_cell(
            table.cell(row_idx, 0),
            str(item["numero"]),
            bg_color="FFFFFF",
            bold=False,
            size=12,
            font_color="000000"
        )

        for sem_idx, semana in enumerate(item["semanas"], start=1):
            start_col = week_starts[sem_idx - 1]
            end_col = start_col + 5
            bloques = semana["bloques"]

            if len(bloques) == 0:
                merged = table.cell(row_idx, start_col)
                merged.merge(table.cell(row_idx, end_col))
                style_word_cell(merged, "", bg_color="FFFFFF", size=12)

            elif len(bloques) == 1:
                b = bloques[0]
                merged = table.cell(row_idx, start_col)
                merged.merge(table.cell(row_idx, end_col))
                style_word_cell(
                    merged,
                    str(b["valor"]),
                    bg_color=b["color"],
                    font_color="FFFFFF",
                    bold=True,
                    size=12
                )

            elif len(bloques) == 2:
                rangos = [(start_col, start_col + 2), (start_col + 3, end_col)]
                for b, (r1, r2) in zip(bloques, rangos):
                    merged = table.cell(row_idx, r1)
                    merged.merge(table.cell(row_idx, r2))
                    style_word_cell(
                        merged,
                        str(b["valor"]),
                        bg_color=b["color"],
                        font_color="FFFFFF",
                        bold=True,
                        size=12
                    )

            else:
                total_week_cols = end_col - start_col + 1
                n = len(bloques)
                ancho_base = total_week_cols // n
                resto = total_week_cols % n

                rangos = []
                actual = start_col
                for j in range(n):
                    ancho = ancho_base + (1 if j < resto else 0)
                    fin = actual + ancho - 1
                    rangos.append((actual, fin))
                    actual = fin + 1

                for b, (r1, r2) in zip(bloques, rangos):
                    merged = table.cell(row_idx, r1)
                    merged.merge(table.cell(row_idx, r2))
                    style_word_cell(
                        merged,
                        str(b["valor"]),
                        bg_color=b["color"],
                        font_color="FFFFFF",
                        bold=True,
                        size=12
                    )

    doc.add_paragraph("")

def construir_word_final(puesto, mes1, anio1, mes2, anio2, rendimiento1, rendimiento2, output_file):
    doc = Document()
    configure_document_landscape(doc)

    add_rendimiento_table_to_doc(
        doc=doc,
        puesto=puesto,
        mes_anio=f"{mes1} {anio1}",
        rendimiento=rendimiento1
    )

    doc.add_page_break()

    add_rendimiento_table_to_doc(
        doc=doc,
        puesto=puesto,
        mes_anio=f"{mes2} {anio2}",
        rendimiento=rendimiento2
    )

    doc.save(output_file)

