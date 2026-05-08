import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from copy import deepcopy
from io import BytesIO
import matplotlib.font_manager as fm
import os

#st.set_page_config(page_title="Generador de Acuse", layout="centered")

RUTA_CONTENIDO = os.path.join("datosp", "7. ACUSE.docx")

def mostrar_modulo_acuse():

    st.title("Generador de Acuse de Recibo")

    plantilla_subida = st.file_uploader(
        "Sube la plantilla Word con encabezado/diseño",
        type=["docx"]
    )

    empresa_recibe = st.text_input("Empresa que recibe el servicio")
    empresa_brinda = st.text_input("Empresa que brinda el servicio")
    nombre_programa = st.text_area("Nombre del programa")

    fuentes = sorted(set([f.name for f in fm.fontManager.ttflist]))

    st.subheader("Configuración de estilos del documento")

    col1, col2 = st.columns(2)

    with col1:
        fuente_titulo1 = st.selectbox("Fuente Título 1", fuentes, key="fuente_titulo1")
        tamano_titulo1 = st.number_input("Tamaño Título 1", min_value=6, max_value=60, value=16)
        color_titulo1 = st.color_picker("Color Título 1", "#000000")

        fuente_titulo2 = st.selectbox("Fuente Título 2", fuentes, key="fuente_titulo2")
        tamano_titulo2 = st.number_input("Tamaño Título 2", min_value=6, max_value=60, value=14)
        color_titulo2 = st.color_picker("Color Título 2", "#000000")

    with col2:
        fuente_titulo3 = st.selectbox("Fuente Título 3", fuentes, key="fuente_titulo3")
        tamano_titulo3 = st.number_input("Tamaño Título 3", min_value=6, max_value=60, value=12)
        color_titulo3 = st.color_picker("Color Título 3", "#000000")

        fuente_normal = st.selectbox("Fuente Texto normal", fuentes, key="fuente_normal")
        tamano_normal = st.number_input("Tamaño Texto normal", min_value=6, max_value=60, value=11)
        color_normal = st.color_picker("Color Texto normal", "#000000")


def hex_a_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def aplicar_fuente(run, fuente, tamano, color_hex):
    run.font.name = fuente
    run.font.size = Pt(tamano)

    r, g, b = hex_a_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts

    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", fuente)
    rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", fuente)
    rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", fuente)
    rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs", fuente)


def obtener_formato_por_estilo(parrafo, configuracion):
    try:
        nombre_estilo = parrafo.style.name
    except:
        nombre_estilo = ""

    if nombre_estilo in ["Título 1", "Heading 1"]:
        return configuracion["titulo1"]

    elif nombre_estilo in ["Título 2", "Heading 2"]:
        return configuracion["titulo2"]

    elif nombre_estilo in ["Título 3", "Heading 3"]:
        return configuracion["titulo3"]

    else:
        return configuracion["normal"]


def reemplazar_en_parrafo(parrafo, reemplazos, configuracion):
    texto = "".join(run.text for run in parrafo.runs)

    if not texto:
        return

    for clave, valor in reemplazos.items():
        texto = texto.replace(clave, valor)

    fuente, tamano, color = obtener_formato_por_estilo(parrafo, configuracion)

    for run in parrafo.runs:
        run.text = ""

    if parrafo.runs:
        parrafo.runs[0].text = texto
        aplicar_fuente(parrafo.runs[0], fuente, tamano, color)
    else:
        run = parrafo.add_run(texto)
        aplicar_fuente(run, fuente, tamano, color)


def procesar_documento(doc, reemplazos, configuracion):
    for parrafo in doc.paragraphs:
        reemplazar_en_parrafo(parrafo, reemplazos, configuracion)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_parrafo(parrafo, reemplazos, configuracion)

    return doc


def limpiar_cuerpo_documento(doc):
    body = doc._body._element

    for elemento in list(body):
        if elemento.tag.endswith("sectPr"):
            continue
        body.remove(elemento)


def insertar_contenido(destino, origen):
    body_destino = destino._body._element
    body_origen = origen._body._element

    sectPr = body_destino.sectPr

    for elemento in body_origen:
        if elemento.tag.endswith("sectPr"):
            continue

        nuevo_elemento = deepcopy(elemento)

        if sectPr is not None:
            body_destino.insert(body_destino.index(sectPr), nuevo_elemento)
        else:
            body_destino.append(nuevo_elemento)

    return destino


if st.button("Generar Word"):
    if plantilla_subida is None:
        st.error("Primero sube la plantilla Word.")

    elif not os.path.exists(RUTA_CONTENIDO):
        st.error(f"No se encontró el archivo: {RUTA_CONTENIDO}")

    elif not empresa_recibe or not empresa_brinda or not nombre_programa:
        st.error("Completa todos los campos.")

    else:
        plantilla = Document(plantilla_subida)
        contenido = Document(RUTA_CONTENIDO)

        configuracion = {
            "titulo1": (fuente_titulo1, tamano_titulo1, color_titulo1),
            "titulo2": (fuente_titulo2, tamano_titulo2, color_titulo2),
            "titulo3": (fuente_titulo3, tamano_titulo3, color_titulo3),
            "normal": (fuente_normal, tamano_normal, color_normal)
        }

        reemplazos = {
            "{{EMPRESA_RECIBE}}": empresa_recibe.upper(),
            "{{EMPRESA_BRINDA}}": empresa_brinda.upper(),
            "{{NOMBRE_PROGRAMA}}": nombre_programa.upper()
        }

        contenido = procesar_documento(
            contenido,
            reemplazos,
            configuracion
        )

        limpiar_cuerpo_documento(plantilla)

        documento_final = insertar_contenido(
            plantilla,
            contenido
        )

        salida = BytesIO()
        documento_final.save(salida)
        salida.seek(0)

        st.success("Documento generado correctamente.")

        st.download_button(
            label="Descargar Word generado",
            data=salida,
            file_name="acuse_generado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )