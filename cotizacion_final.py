import re
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
import random
import math
import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from datetime import timedelta, datetime
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from utils.calendario_mexico import calcular_fecha_larga_habil_mexico
from utils.moneda import formatear_moneda, convertir_numero_a_letras_mxn
from utils.word_table import (
    poner_bordes_tabla,
    colorear_celda,
    configurar_texto_celda
)
from utils.word_styles import aplicar_heading_a_texto

# Ajusta estos imports según tu proyecto
from app.db.database import SessionLocal
from app.models.empresa import Empresa
from app.models.empresa_plantilla_word import EmpresaPlantillaWord
from app.models.empresa_estilo_word import EmpresaEstiloWord
from app.models.empresa_plantilla_asignacion import EmpresaPlantillaAsignacion

load_dotenv()

client = OpenAI(
    api_key=os.environ["OPENAI_API_KEY"]
)

ARCHIVO_ACUSE = "7.ACUSE.docx"
ARCHIVO_RESUMEN = "6.RESUMEN-EJECUTIVO.docx"
ARCHIVO_CALENDARIO = "4.CALENDARIO-DE-TRABAJO.docx"
ARCHIVO_COTIZACION_INICIAL = "2.COTIZACION-INICIAL.docx"
ARCHIVO_PROPUESTA = "1.PROPUESTA.docx"
ARCHIVO_ENTREGABLE = "5.ENTREGABLE.docx"


def llamar_openai_texto(prompt):
    respuesta = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )

    return respuesta.output_text.strip()


def generar_objetivo_general(nombre_servicio):
    prompt = f"""
Actúa como consultor corporativo experto en redacción de propuestas comerciales empresariales.

Redacta un párrafo formal, profesional y persuasivo para una propuesta de servicios especializada en: “{nombre_servicio}”.

El texto debe iniciar con la frase:

Es un placer presentarle la siguiente propuesta de servicios especializados en “{nombre_servicio.upper()}”, cuyo objetivo es...

No uses viñetas ni subtítulos. Solo entrega el párrafo final.
"""
    return llamar_openai_texto(prompt)


def generar_introduccion(nombre_servicio, conceptos_factura):
    prompt = f"""
Actúa como consultor corporativo senior especializado en desarrollo empresarial, consultoría estratégica y redacción ejecutiva de propuestas profesionales para empresas en México.

Tu tarea es redactar exclusivamente el apartado:

“INTRODUCCIÓN”

con base en la siguiente información:

TEMA PRINCIPAL DEL SERVICIO:
{nombre_servicio}

CONCEPTOS RELACIONADOS:
{conceptos_factura}

OBJETIVO DE LA REDACCIÓN:

Debes analizar e interpretar tanto la temática principal como los conceptos relacionados para comprender el alcance real del servicio prestado y transformarlo en una introducción ejecutiva, elegante y estratégicamente estructurada.

NO debes copiar literalmente los conceptos proporcionados.
NO debes describirlos de forma mecánica.
Debes convertir la información en una narrativa corporativa profesional con sentido empresarial y consultivo.

INSTRUCCIONES OBLIGATORIAS:

1. La introducción debe:
- Explicar la importancia estratégica del tema dentro de una organización.
- Justificar la necesidad del servicio.
- Contextualizar problemáticas, retos u oportunidades empresariales.
- Introducir el propósito general del proyecto.
- Reflejar impacto organizacional y operativo.
- Transmitir valor empresarial y visión estratégica.

2. El texto debe mencionar de forma natural beneficios relacionados con:
- Optimización de procesos.
- Eficiencia operativa.
- Fortalecimiento organizacional.
- Crecimiento empresarial.
- Control administrativo.
- Rentabilidad.
- Toma de decisiones.
- Posicionamiento.
- Productividad.
- Competitividad.

3. El estilo de redacción debe ser:
- Corporativo.
- Ejecutivo.
- Técnico.
- Elegante.
- Consultivo.
- Profesional.
- Propio de una firma de consultoría empresarial premium.

4. La redacción debe:
- Tener alta coherencia.
- Sonar natural y sofisticada.
- Evitar redundancias.
- Evitar frases vacías o genéricas.
- Evitar lenguaje informal.
- Mantener profundidad estratégica y empresarial.

5. FORMATO OBLIGATORIO:
- Un solo párrafo continuo.
- Sin subtítulos.
- Sin listas.
- Sin viñetas.
- Sin saltos de línea.
- Sin encabezados adicionales.

6. El resultado debe sentirse como:
- Una introducción ejecutiva de alto nivel.
- Parte de una propuesta formal empresarial.
- Un documento consultivo premium.
- Un texto con alto valor percibido.

7. NO expliques el proceso.
8. NO uses frases repetitivas.
9. SOLO entrega el párrafo final listo para integrarse en una propuesta formal.

Genera una introducción amplia, sólida, estratégica y con profundidad corporativa.
"""

    return llamar_openai_texto(prompt)


def generar_problematica(nombre_servicio, conceptos_factura):
    prompt = f"""
Actúa como consultor corporativo senior especializado en diagnóstico empresarial, análisis de riesgos, desarrollo organizacional y redacción estratégica de propuestas profesionales para empresas en México.

Tu tarea es redactar exclusivamente el apartado:

“PROBLEMÁTICA”

con base en la siguiente información:

TEMA PRINCIPAL DEL SERVICIO:
{nombre_servicio}

CONCEPTOS RELACIONADOS:
{conceptos_factura}

OBJETIVO DE LA REDACCIÓN:

Debes analizar e interpretar la temática principal y los conceptos relacionados para identificar riesgos, deficiencias, necesidades, brechas operativas o áreas de oportunidad que una empresa podría enfrentar cuando no cuenta con una solución especializada en esta materia.

NO debes copiar literalmente los conceptos proporcionados.
NO debes redactar una descripción genérica del servicio.
Debes convertir la información en una problemática empresarial sólida, realista y estratégicamente justificada.

INSTRUCCIONES OBLIGATORIAS:

1. La problemática debe:
- Describir las principales deficiencias, limitaciones o necesidades que puede enfrentar una empresa en relación con el servicio indicado.
- Exponer consecuencias operativas, administrativas, comerciales, financieras, organizacionales o estratégicas.
- Evidenciar riesgos por falta de planeación, control, seguimiento, estructura, eficiencia, cumplimiento, coordinación o especialización.
- Justificar la necesidad de apoyo profesional especializado.
- Cerrar con una idea que refuerce la importancia de implementar soluciones alineadas con la temática del servicio.

2. El texto debe abordar, cuando sea pertinente, impactos como:
- Ineficiencia operativa.
- Falta de control administrativo.
- Pérdida de oportunidades comerciales.
- Incremento de costos.
- Desorganización interna.
- Debilidad en la toma de decisiones.
- Riesgos de cumplimiento.
- Baja productividad.
- Limitaciones para el crecimiento.
- Pérdida de competitividad.
- Falta de trazabilidad o seguimiento.

3. El estilo de redacción debe ser:
- Corporativo.
- Técnico.
- Formal.
- Ejecutivo.
- Analítico.
- Consultivo.
- Propio de una firma de consultoría empresarial premium.

4. La redacción debe:
- Tener profundidad empresarial.
- Sonar natural, seria y profesional.
- Evitar exageraciones poco creíbles.
- Evitar frases repetitivas.
- Evitar lenguaje informal.
- Evitar copiar los conceptos de factura.
- Mantener coherencia con la temática del servicio.

5. FORMATO OBLIGATORIO:
- Un solo párrafo continuo.
- Sin subtítulos.
- Sin listas.
- Sin viñetas.
- Sin saltos de línea.
- Sin encabezados adicionales.

6. El resultado final debe sentirse como:
- Un diagnóstico empresarial de alto nivel.
- Una problemática clara y bien justificada.
- Parte de una propuesta formal de servicios.
- Un texto con alto valor consultivo y estratégico.

7. NO expliques el proceso.
8. NO menciones que analizaste los conceptos.
9. NO uses frases como “la problemática es”.
10. SOLO entrega el párrafo final listo para integrarse en una propuesta formal.

Genera una problemática amplia, sólida, realista y con enfoque corporativo.
"""

    return llamar_openai_texto(prompt)


def generar_objetivo_general_tabla(nombre_servicio, conceptos_factura, nombre_cliente):
    prompt = f"""
Actúa como consultor corporativo senior especializado en desarrollo empresarial, planeación estratégica y redacción ejecutiva de propuestas profesionales para empresas en México.

Tu tarea es redactar exclusivamente el apartado:

“OBJETIVO GENERAL”

con base en la siguiente información:

TEMA PRINCIPAL DEL SERVICIO:
{nombre_servicio}

CONCEPTOS RELACIONADOS:
{conceptos_factura}

CLIENTE:
{nombre_cliente}

OBJETIVO DE LA REDACCIÓN:

Debes analizar e interpretar la temática principal y los conceptos relacionados para comprender el propósito estratégico del servicio, identificando los resultados, beneficios, mejoras y alcances empresariales que la organización obtendrá al finalizar el proyecto.

NO debes copiar literalmente los conceptos proporcionados.
NO debes generar texto genérico.
Debes transformar la información en una redacción ejecutiva, sólida y estratégicamente estructurada.

INSTRUCCIONES OBLIGATORIAS:

1. El texto DEBE iniciar exactamente con:

“Al finalizar el servicio, {nombre_cliente.upper()}, será capaz de...”

2. Después del inicio, desarrolla de forma natural y elegante:

- El resultado global que obtendrá la empresa.
- El propósito estratégico del servicio.
- Las capacidades, mejoras o fortalecimientos organizacionales alcanzados.
- Beneficios operativos, administrativos, comerciales, financieros o estratégicos.
- Cómo el servicio contribuirá al crecimiento, eficiencia, control, posicionamiento, competitividad, rentabilidad o reducción de riesgos.
- El impacto positivo en la toma de decisiones, estructura organizacional y desempeño empresarial.

3. El texto debe transmitir:
- Alto valor percibido.
- Nivel corporativo premium.
- Solidez estratégica y técnica.
- Formalidad ejecutiva.
- Visión empresarial y consultiva.

4. El estilo de redacción debe ser:
- Corporativo.
- Ejecutivo.
- Técnico.
- Elegante.
- Consultivo.
- Persuasivo.
- Propio de una firma consultora empresarial de alto nivel.

5. La redacción debe:
- Sonar natural, profesional y sofisticada.
- Evitar redundancias.
- Evitar frases vacías o genéricas.
- Evitar lenguaje informal.
- Mantener coherencia empresarial y estratégica.
- Tener profundidad consultiva y organizacional.

6. FORMATO OBLIGATORIO:
- Un solo párrafo amplio y continuo.
- Sin subtítulos.
- Sin listas.
- Sin viñetas.
- Sin saltos de línea.
- Sin encabezados adicionales.

7. El resultado final debe sentirse como:
- Un objetivo general ejecutivo de alto nivel.
- Parte de una propuesta formal empresarial.
- Un documento consultivo premium.
- Un texto con enfoque estratégico y corporativo.

8. NO expliques el proceso.
9. NO repitas literalmente los conceptos proporcionados.
10. SOLO entrega el párrafo final listo para integrarse en una propuesta formal.

Genera un objetivo general amplio, sólido, estratégico y con profundidad corporativa.
"""

    return llamar_openai_texto(prompt)


def generar_objetivos_especificos(nombre_servicio, conceptos_factura, nombre_cliente):
    prompt = f"""
Actúa como consultor corporativo senior especializado en desarrollo empresarial, planeación estratégica, optimización organizacional y redacción ejecutiva de propuestas profesionales para empresas en México.

Tu tarea es redactar exclusivamente los:

“OBJETIVOS ESPECÍFICOS”

con base en la siguiente información:

TEMA PRINCIPAL DEL SERVICIO:
{nombre_servicio}

CONCEPTOS RELACIONADOS:
{conceptos_factura}

CLIENTE:
{nombre_cliente}

OBJETIVO DE LA REDACCIÓN:

Debes analizar e interpretar la temática principal y los conceptos relacionados para identificar acciones estratégicas, mejoras operativas, beneficios organizacionales y resultados empresariales que serán alcanzados mediante el servicio prestado.

NO debes copiar literalmente los conceptos proporcionados.
NO debes redactar objetivos genéricos.
Debes transformar la información en objetivos específicos ejecutivos, técnicos y estratégicamente estructurados.

INSTRUCCIONES OBLIGATORIAS:

1. Genera entre 4 y 6 objetivos específicos.

2. TODOS los objetivos deben iniciar EXACTAMENTE con:

“{nombre_cliente.upper()} al término del servicio,...”

3. Cada objetivo debe expresar de forma clara y profesional:
- Resultados concretos.
- Mejoras empresariales.
- Fortalecimiento organizacional.
- Optimización operativa.
- Incremento de eficiencia.
- Implementación de estrategias.
- Desarrollo de capacidades.
- Control y seguimiento.
- Reducción de riesgos.
- Mejora en toma de decisiones.
- Crecimiento o consolidación empresarial.

4. Utiliza de forma natural verbos estratégicos y consultivos como:
- Diagnosticar
- Diseñar
- Implementar
- Optimizar
- Fortalecer
- Mejorar
- Estructurar
- Incrementar
- Consolidar
- Evaluar
- Medir
- Desarrollar
- Integrar
- Supervisar
- Coordinar

5. El estilo de redacción debe ser:
- Corporativo.
- Ejecutivo.
- Técnico.
- Formal.
- Consultivo.
- Estratégico.
- Propio de una firma consultora empresarial premium.

6. La redacción debe:
- Sonar sofisticada y profesional.
- Evitar redundancias.
- Evitar frases vacías.
- Evitar lenguaje informal.
- Tener coherencia empresarial y estratégica.
- Mantener profundidad organizacional y consultiva.

7. FORMATO OBLIGATORIO:

1. Texto...
2. Texto...
3. Texto...
4. Texto...

8. NO agregues títulos.
9. NO agregues introducciones.
10. NO agregues explicaciones adicionales.
11. SOLO entrega el resultado final.
12. Cada objetivo debe ser amplio, sólido y con alto valor percibido.

Genera objetivos específicos estratégicos, ejecutivos y corporativamente robustos.
"""

    return llamar_openai_texto(prompt)


def generar_metodologia(nombre_servicio, conceptos_factura):
    prompt = f"""
Actúa como consultor corporativo senior especializado en desarrollo empresarial, diseño metodológico, optimización organizacional y redacción estratégica de propuestas profesionales para empresas en México.

Tu tarea es redactar exclusivamente el apartado:

“METODOLOGÍA”

con base en la siguiente información:

TEMA PRINCIPAL DEL SERVICIO:
{nombre_servicio}

CONCEPTOS RELACIONADOS:
{conceptos_factura}

OBJETIVO DE LA REDACCIÓN:

Debes interpretar los conceptos relacionados como componentes metodológicos del servicio, transformándolos en temas técnicos y estratégicos desarrollados profesionalmente dentro de una metodología empresarial formal.

NO debes copiar únicamente los conceptos.
NO debes limitarte a describirlos superficialmente.
Debes convertir cada concepto en un apartado metodológico sólido, ejecutivo y consultivamente estructurado.

INSTRUCCIONES OBLIGATORIAS:

1. La metodología debe construirse directamente a partir de los conceptos relacionados.

2. El número total de apartados metodológicos generados DEBE ser exactamente igual al número de conceptos proporcionados.

3. Cada concepto debe convertirse obligatoriamente en un único apartado metodológico individual.

4. Está estrictamente prohibido:
- Fusionar conceptos.
- Eliminar conceptos.
- Omitir conceptos.
- Reemplazar conceptos por otros nombres.
- Cambiar el nombre original de los conceptos.
- Resumir varios conceptos en uno solo.

5. Los nombres de los subtítulos deben conservar EXACTAMENTE el mismo texto del concepto original proporcionado, únicamente convirtiéndolo a MAYÚSCULAS y respetando acentos.

6. Cada concepto debe transformarse en:
- Un subtítulo metodológico profesional conservando exactamente el nombre original.
- Un desarrollo técnico y corporativo que explique:
    - Qué se analiza.
    - Qué se desarrolla.
    - Qué se implementa.
    - Qué se optimiza.
    - Qué se supervisa.
    - Qué beneficios empresariales genera.
    - Qué impacto operativo, administrativo, estratégico o comercial aporta.

7. Puedes únicamente:
- Mejorar la redacción técnica y ejecutiva del desarrollo descriptivo.
- Ajustar la secuencia metodológica sin alterar los conceptos.
- Ampliar el alcance técnico y corporativo de cada apartado.

8. Cada explicación debe:
- Tener profundidad empresarial.
- Sonar consultiva y estratégica.
- Mantener coherencia con la temática principal.
- Reflejar valor profesional y organizacional.
- Explicar el alcance funcional del tema.

9. El estilo de redacción debe ser:
- Corporativo.
- Ejecutivo.
- Técnico.
- Elegante.
- Consultivo.
- Profesional.
- Propio de una firma consultora empresarial premium.

10. La redacción debe:
- Evitar frases genéricas.
- Evitar redundancias.
- Evitar explicaciones simples o pobres.
- Evitar lenguaje informal.
- Tener claridad técnica y estratégica.
- Mantener alto valor percibido.

FORMATO OBLIGATORIO:

Primero escribe EXACTAMENTE:

A continuación, se presentan los temas incluidos en el servicio:

Después desarrolla cada apartado usando EXACTAMENTE esta estructura:

[SUBTÍTULO EN MAYÚSCULAS]. Desarrollo profesional del tema.

REGLAS ESTRICTAS DE FORMATO:

- TODOS los subtítulos deben estar completamente en MAYÚSCULAS.
- TODOS los subtítulos deben conservar exactamente el nombre original del concepto.
- TODOS los subtítulos deben ir en negritas.
- TODOS los subtítulos deben terminar en punto dentro de las negritas.
- Respeta tildes aun estando en mayúsculas.
- NO uses viñetas.
- NO uses numeración.
- Cada apartado debe estar separado por un salto de línea.
- NO agregues títulos adicionales.
- NO agregues introducciones fuera de la frase solicitada.
- NO agregues conclusiones.
- SOLO entrega el resultado final.
- Verifica obligatoriamente que el número de apartados generados coincida exactamente con el número de conceptos proporcionados.

EJEMPLO DE ESTRUCTURA OBLIGATORIA:

PLAN ESTRATÉGICO PARA LA OPTIMIZACIÓN OPERATIVA Y ADMINISTRATIVA. Desarrollo de mecanismos de análisis, organización y control orientados al fortalecimiento de los procesos internos, permitiendo mejorar la eficiencia operativa, la coordinación organizacional y la toma de decisiones dentro de la empresa.

IMPLEMENTACIÓN DE MECANISMOS DE SUPERVISIÓN Y SEGUIMIENTO OPERATIVO. Integración de herramientas de monitoreo y evaluación continua para dar seguimiento al desempeño de las actividades estratégicas, facilitando el control de resultados, la detección de desviaciones y la mejora continua de la operación.

Genera una metodología amplia, sólida, técnica y corporativamente robusta.
"""

    return llamar_openai_texto(prompt)


def generar_texto_concepto_entregable(nombre_programa, concepto):
    prompt = f"""
Actúa como consultor corporativo senior especializado en redacción técnica, desarrollo metodológico y elaboración de entregables empresariales de alto nivel.

Tu tarea consiste en explicar y desarrollar profesionalmente un concepto específico perteneciente a un programa principal de consultoría o servicio empresarial.

INFORMACIÓN BASE:

NOMBRE DEL PROGRAMA PRINCIPAL:
{nombre_programa}

CONCEPTO O SUBTEMA A DESARROLLAR:
{concepto}

OBJETIVO:

Debes tomar el concepto proporcionado y desarrollar una explicación profesional, amplia y enriquecida sobre lo que representa dentro del contexto del programa principal.

El contenido debe explicar el concepto de forma corporativa, técnica y estratégica, como si fuera un apartado especializado dentro de un entregable ejecutivo o metodología empresarial.

INSTRUCCIONES OBLIGATORIAS:

- Explica claramente qué es el concepto.
- Desarrolla profesionalmente cómo funciona o en qué consiste.
- Explica qué aspectos contempla, analiza, desarrolla u optimiza.
- Relaciona naturalmente el concepto con el programa principal.
- Explica la utilidad empresarial del concepto.
- Describe el valor operativo, administrativo, estratégico u organizacional que aporta.
- El contenido debe sentirse técnico, corporativo y consultivo.
- Enriquece ampliamente la información aunque el concepto original sea corto.
- No copies únicamente el concepto.
- No hagas definiciones simples tipo diccionario.
- No uses frases vacías o genéricas.
- No repitas estructuras innecesarias.
- No uses listas.
- No uses viñetas.
- No uses subtítulos.
- No uses encabezados.
- No menciones que eres IA.
- Evita frases robóticas o demasiado genéricas.
- El estilo debe parecer redactado por una firma consultora empresarial premium.

FORMATO OBLIGATORIO:

- Redacta exactamente 2 párrafos amplios.
- El primer párrafo debe enfocarse en explicar qué es el concepto y qué comprende.
- El segundo párrafo debe enfocarse en su aplicación, utilidad, beneficios e impacto empresarial.
- La redacción debe ser fluida, natural y altamente profesional.

IMPORTANTE:

El resultado debe sentirse como la explicación técnica y estratégica de un subtema especializado derivado de un programa corporativo principal.
"""

    return llamar_openai_texto(prompt)


def generar_introduccion_entregable(
    nombre_programa,
    nombre_cliente,
    empresa_nombre
):
    prompt = f"""
Redacta una introducción corporativa, formal y profesional para un entregable de prestación de servicios.

Cliente:
{nombre_cliente}

Empresa prestadora:
{empresa_nombre}

Servicio:
{nombre_programa}

Instrucciones:
- Redacta entre 3 y 4 párrafos.
- Usa lenguaje empresarial y ejecutivo.
- Explica la importancia del servicio realizado.
- Explica el valor operativo y administrativo del servicio.
- No uses listas.
- No uses subtítulos.
- No menciones inteligencia artificial.
- El tono debe ser corporativo y técnico.
"""

    return llamar_openai_texto(prompt)


def generar_referencias_entregable(nombre_programa, conceptos):
    anio_actual = datetime.now().year
    anio_minimo = anio_actual - 5
    total_referencias = len(conceptos)

    conceptos_texto = "\n".join(
        f"{i + 1}. {concepto}"
        for i, concepto in enumerate(conceptos)
    )

    prompt = f"""
Genera referencias bibliográficas en formato APA 7 en español, relacionadas con el siguiente programa:

{nombre_programa}

Conceptos del servicio:
{conceptos_texto}

Instrucciones obligatorias:
- Genera exactamente {total_referencias} referencias.
- Debe haber una referencia relacionada con cada concepto del servicio.
- Las referencias deben estar en español.
- Las referencias deben ser aplicables al contexto empresarial de México.
- Las referencias deben tener una antigüedad máxima de 5 años.
- El año actual es {anio_actual}.
- Solo puedes usar referencias publicadas entre {anio_minimo} y {anio_actual}.
- No uses referencias anteriores a {anio_minimo}.
- Usa formato APA 7.
- Solo entrega la lista de referencias.
- No expliques nada.
"""

    return llamar_openai_texto(prompt)


def limpiar_texto(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()

def obtener_estilos_word_por_plantilla(plantilla_id):
    db = SessionLocal()

    try:
        estilos = (
            db.query(EmpresaEstiloWord)
            .filter(EmpresaEstiloWord.plantilla_id == plantilla_id)
            .all()
        )

        return {
            estilo.clave_estilo: {
                "tipografia": estilo.tipografia,
                "tamanio_letra": estilo.tamanio_letra,
                "color_letra": estilo.color_letra,
                "color_fondo": estilo.color_fondo,
                "negrita": estilo.negrita,
                "cursiva": estilo.cursiva,
                "alineacion": estilo.alineacion,
            }
            for estilo in estilos
        }

    finally:
        db.close()

def convertir_color_bd_a_rgb(color):
    """
    Convierte colores guardados en BD como:
    '#D84B20'
    'D84B20'
    '216,75,32'
    a tupla RGB: (216, 75, 32)
    """
    if not color:
        return (0, 0, 0)

    color = str(color).strip()

    if "," in color:
        partes = color.split(",")
        return tuple(int(p.strip()) for p in partes)

    color = color.replace("#", "")

    if len(color) != 6:
        raise ValueError(f"Color inválido en BD: {color}")

    return tuple(int(color[i:i + 2], 16) for i in (0, 2, 4))


def extraer_entre_parentesis(texto):
    texto = limpiar_texto(texto)
    encontrados = re.findall(r"\((.*?)\)", texto)
    if encontrados:
        return encontrados[-1].strip()
    return texto


def generar_incremento_aleatorio(minimo=5000, maximo=30000):
    return random.randint(minimo, maximo)


def redondear_monto_inicial(monto):
    """
    Convierte el monto a entero sin centavos.
    Ejemplo: 29153.12 + random = 35500
    """
    return math.ceil(monto)


def crear_registros_cotizacion_inicial(registros):
    registros_iniciales = []

    for r in registros:
        monto_final = float(r["Monto"])
        incremento = generar_incremento_aleatorio(5000, 30000)
        monto_inicial = redondear_monto_inicial(monto_final + incremento)

        registros_iniciales.append({
            "Concepto": r["Concepto"],
            "Monto": monto_inicial,
            "Fecha": r["Fecha"],
        })

    return registros_iniciales


def extraer_datos_cotizacion_excel(archivo_excel):
    df = pd.read_excel(archivo_excel, header=None)

    fila_encabezados = None
    for i in range(len(df)):
        fila = [limpiar_texto(x).lower() for x in df.iloc[i].tolist()]
        if "concepto" in fila and "total factura ($)" in fila and "fecha" in fila:
            fila_encabezados = i
            break

    if fila_encabezados is None:
        raise ValueError(
            "No se encontraron los encabezados 'Concepto', "
            "'Total Factura ($)' y 'Fecha' en el Excel."
        )

    encabezados = [limpiar_texto(x) for x in df.iloc[fila_encabezados].tolist()]

    col_concepto = encabezados.index("Concepto")
    col_monto = encabezados.index("Total Factura ($)")
    col_fecha = encabezados.index("Fecha")

    registros = []
    ultimo_concepto = ""

    for i in range(fila_encabezados + 1, len(df)):
        concepto_raw = limpiar_texto(df.iloc[i, col_concepto])
        monto = df.iloc[i, col_monto]
        fecha = df.iloc[i, col_fecha]

        if concepto_raw:
            ultimo_concepto = extraer_entre_parentesis(concepto_raw)

        if pd.notna(monto):
            fecha_convertida = pd.to_datetime(fecha, errors="coerce", dayfirst=True)

            fecha_formateada = ""
            if pd.notna(fecha_convertida):
                fecha_formateada = fecha_convertida.strftime("%d/%m/%Y")

            registros.append({
                "Concepto": ultimo_concepto,
                "Monto": float(monto),
                "Fecha": fecha_formateada,
            })

    if not registros:
        raise ValueError("No se encontraron registros con monto en el Excel.")

    return registros


def extraer_datos_calendario_excel(archivo_excel):
    archivo_excel.seek(0)

    df = pd.read_excel(archivo_excel, header=None)

    fila_encabezados = None

    for i in range(len(df)):
        fila = [limpiar_texto(x).lower() for x in df.iloc[i].tolist()]

        if "concepto" in fila and "fecha" in fila:
            fila_encabezados = i
            break

    if fila_encabezados is None:
        raise ValueError(
            "No se encontraron los encabezados 'Concepto' y 'Fecha' en el Excel."
        )

    encabezados = [
        limpiar_texto(x)
        for x in df.iloc[fila_encabezados].tolist()
    ]

    col_concepto = encabezados.index("Concepto")
    col_fecha = encabezados.index("Fecha")

    bloques = []
    bloque_actual = None

    for i in range(fila_encabezados + 1, len(df)):
        concepto_raw = limpiar_texto(df.iloc[i, col_concepto])
        fecha_raw = df.iloc[i, col_fecha]

        if concepto_raw:
            bloque_actual = {
                "concepto": extraer_entre_parentesis(concepto_raw),
                "fechas": [],
            }

            bloques.append(bloque_actual)

        if bloque_actual is not None:
            fecha = pd.to_datetime(
                fecha_raw,
                errors="coerce",
                dayfirst=True
            )

            if pd.notna(fecha):
                bloque_actual["fechas"].append(fecha.date())

    bloques_validos = []

    for bloque in bloques:
        if bloque["fechas"]:
            bloques_validos.append(bloque)

    for i, bloque in enumerate(bloques_validos):
        fecha_inicio = min(bloque["fechas"])

        if i < len(bloques_validos) - 1:
            siguiente_inicio = min(bloques_validos[i + 1]["fechas"])

            if siguiente_inicio == fecha_inicio:
                fecha_fin = fecha_inicio
            else:
                fecha_fin = siguiente_inicio - timedelta(days=1)
        else:
            fecha_fin = max(bloque["fechas"])

        bloque["fecha_inicio"] = fecha_inicio
        bloque["fecha_fin"] = fecha_fin

    return bloques_validos


def agrupar_conceptos_por_monto(registros):
    grupos = {}

    for r in registros:
        concepto = r["Concepto"]
        monto = r["Monto"]

        if concepto not in grupos:
            grupos[concepto] = []

        grupos[concepto].append(monto)

    return grupos


def obtener_fecha_mas_antigua_registros(registros):
    fechas = []

    for r in registros:
        fecha = pd.to_datetime(r.get("Fecha"), errors="coerce", dayfirst=True)

        if pd.notna(fecha):
            fechas.append(fecha)

    if not fechas:
        raise ValueError("No se encontraron fechas válidas en el Excel.")

    fecha_mas_antigua = min(fechas)

    return fecha_mas_antigua.strftime("%d/%m/%Y")


def calcular_fechas_entregables(registros):
    fecha_mas_antigua = obtener_fecha_mas_antigua_registros(registros)

    fecha_propuesta = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=0
    )

    fecha_cotizacion_inicial = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=2
    )

    fecha_cotizacion_final = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=4
    )

    fecha_calendario = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=6
    )

    return {
        "propuesta": f"Ciudad de Mérida, Yucatán a {fecha_propuesta}",
        "cotizacion_inicial": f"Mérida, Yucatán, México a {fecha_cotizacion_inicial}",
        "cotizacion_final": f"Mérida, Yucatán, México a {fecha_cotizacion_final}",
        "calendario": f"Mérida, Yucatán, México a {fecha_calendario}",
    }


def obtener_empresas_con_plantilla():
    db = SessionLocal()

    try:
        resultados = (
            db.query(Empresa, EmpresaPlantillaWord, EmpresaPlantillaAsignacion)
            .join(
                EmpresaPlantillaAsignacion,
                EmpresaPlantillaAsignacion.empresa_externa_id == Empresa.id
            )
            .join(
                EmpresaPlantillaWord,
                EmpresaPlantillaWord.id == EmpresaPlantillaAsignacion.plantilla_id
            )
            .filter(EmpresaPlantillaAsignacion.activo == True)
            .filter(EmpresaPlantillaAsignacion.membrete_path.isnot(None))
            .filter(EmpresaPlantillaWord.plantilla_path.isnot(None))
            .order_by(Empresa.nombre.asc())
            .all()
        )

        empresas = {}

        for empresa, plantilla, asignacion in resultados:
            empresas[empresa.nombre] = {
                "empresa_id": empresa.id,
                "plantilla_id": plantilla.id,
                "nombre": empresa.nombre,
                "razon_social": empresa.razon_social,
                "membrete_path": asignacion.membrete_path,
                "plantilla_path": plantilla.plantilla_path,
                "color_texto_base": plantilla.color_texto_base,
                "color_primario": plantilla.color_primario,
                "color_secundario": plantilla.color_secundario,
                "tipografia_base": plantilla.tipografia_base,
                "tamanio_base": plantilla.tamanio_base,
            }

        return empresas

    finally:
        db.close()


def reemplazar_texto_en_parrafo(parrafo, reemplazos):
    texto_original = parrafo.text

    texto_nuevo = texto_original

    for clave, valor in reemplazos.items():
        texto_nuevo = texto_nuevo.replace(clave, str(valor))

    if texto_nuevo == texto_original:
        return

    # borrar runs existentes
    for run in parrafo.runs:
        run.text = ""

    # crear nuevo run limpio
    nuevo_run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()

    nuevo_run.text = texto_nuevo


def reemplazar_parametros_documento(doc, reemplazos):
    for parrafo in doc.paragraphs:
        reemplazar_texto_en_parrafo(parrafo, reemplazos)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_texto_en_parrafo(parrafo, reemplazos)


def aplicar_negrita_a_texto(doc, texto_objetivo, estilo_base=None):
    if not texto_objetivo:
        return

    texto_objetivo = str(texto_objetivo).strip()

    def aplicar_formato_run(run, negrita=False):
        if estilo_base:
            run.font.name = estilo_base["tipografia"]
            run.font.size = Pt(int(estilo_base["tamanio_letra"]))
            run.font.color.rgb = RGBColor(
                *convertir_color_bd_a_rgb(estilo_base["color_letra"])
            )
            run.font.italic = bool(estilo_base["cursiva"])

        run.font.bold = negrita

    def procesar_parrafo(parrafo):
        texto_completo = "".join(run.text for run in parrafo.runs)

        if texto_objetivo not in texto_completo:
            return

        for run in parrafo.runs:
            run.text = ""

        partes = texto_completo.split(texto_objetivo)

        for i, parte in enumerate(partes):
            if parte:
                run_normal = parrafo.add_run(parte)
                aplicar_formato_run(run_normal, negrita=False)

            if i < len(partes) - 1:
                run_negrita = parrafo.add_run(texto_objetivo)
                aplicar_formato_run(run_negrita, negrita=True)

    for parrafo in doc.paragraphs:
        procesar_parrafo(parrafo)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    procesar_parrafo(parrafo)


def insertar_tabla_despues_de_parrafo(parrafo, tabla):
    parrafo._p.addnext(tabla._tbl)


def eliminar_parrafo(parrafo):
    elemento = parrafo._element
    elemento.getparent().remove(elemento)
    parrafo._p = parrafo._element = None


def buscar_parrafo_con_texto(doc, texto_buscado):
    for parrafo in doc.paragraphs:
        if texto_buscado in parrafo.text:
            return parrafo

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    if texto_buscado in parrafo.text:
                        return parrafo

    return None


def crear_tabla_cotizacion(doc, registros, fuente, size_letra, paleta):
    grupos = agrupar_conceptos_por_monto(registros)

    tabla = doc.add_table(rows=1, cols=2)
    poner_bordes_tabla(tabla)

    header_row = tabla.rows[0]

    configurar_texto_celda(
        header_row.cells[0],
        "CONCEPTOS",
        fuente,
        size_letra,
        True,
        paleta["texto"],
        WD_ALIGN_PARAGRAPH.LEFT,
    )

    configurar_texto_celda(
        header_row.cells[1],
        "MONTO",
        fuente,
        size_letra,
        True,
        paleta["texto"],
        WD_ALIGN_PARAGRAPH.RIGHT,
    )

    colorear_celda(header_row.cells[0], paleta["principal"])
    colorear_celda(header_row.cells[1], paleta["principal"])

    header_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    header_row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    total_general = 0
    indice_concepto = 0

    for concepto, montos in grupos.items():
        filas_concepto = []
        aplicar_relleno = indice_concepto % 2 == 0

        for idx, monto in enumerate(montos):
            row = tabla.add_row()
            filas_concepto.append(row)

            if idx == 0:
                configurar_texto_celda(
                    row.cells[0],
                    concepto.upper(),
                    fuente,
                    size_letra,
                    False,
                    paleta["texto"],
                    WD_ALIGN_PARAGRAPH.LEFT,
                )
            else:
                configurar_texto_celda(
                    row.cells[0],
                    "",
                    fuente,
                    size_letra,
                    False,
                    paleta["texto"],
                    WD_ALIGN_PARAGRAPH.LEFT,
                )

            configurar_texto_celda(
                row.cells[1],
                formatear_moneda(monto),
                fuente,
                size_letra,
                False,
                paleta["texto"],
                WD_ALIGN_PARAGRAPH.RIGHT,
            )

            if aplicar_relleno:
                colorear_celda(row.cells[0], paleta["secundario"])
                colorear_celda(row.cells[1], paleta["secundario"])

            total_general += monto

        if len(filas_concepto) > 1:
            celda_inicio = filas_concepto[0].cells[0]
            celda_fin = filas_concepto[-1].cells[0]
            celda_inicio.merge(celda_fin)

            configurar_texto_celda(
                celda_inicio,
                concepto.upper(),
                fuente,
                size_letra,
                False,
                paleta["texto"],
                WD_ALIGN_PARAGRAPH.LEFT,
            )

            if aplicar_relleno:
                colorear_celda(celda_inicio, paleta["secundario"])

        indice_concepto += 1

    total_row = tabla.add_row()

    configurar_texto_celda(
        total_row.cells[0],
        "TOTAL",
        fuente,
        size_letra,
        True,
        paleta["texto"],
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    configurar_texto_celda(
        total_row.cells[1],
        formatear_moneda(total_general),
        fuente,
        size_letra,
        True,
        paleta["texto"],
        WD_ALIGN_PARAGRAPH.RIGHT,
    )

    letras_row = tabla.add_row()
    celda_letras = letras_row.cells[0].merge(letras_row.cells[1])

    configurar_texto_celda(
        celda_letras,
        convertir_numero_a_letras_mxn(total_general),
        fuente,
        size_letra,
        True,
        paleta["texto"],
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    return tabla, total_general


def aplicar_estilo_base_a_todo(doc, estilo_base):
    fuente = estilo_base["tipografia"]
    size_letra = int(estilo_base["tamanio_letra"])
    color_texto = convertir_color_bd_a_rgb(estilo_base["color_letra"])

    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            run.font.name = fuente
            run.font.size = Pt(size_letra)
            run.font.bold = bool(estilo_base["negrita"])
            run.font.italic = bool(estilo_base["cursiva"])
            run.font.color.rgb = RGBColor(*color_texto)

        if estilo_base["alineacion"]:
            aplicar_alineacion_parrafo(parrafo, estilo_base["alineacion"])

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.name = fuente
                        run.font.size = Pt(size_letra)
                        run.font.bold = bool(estilo_base["negrita"])
                        run.font.italic = bool(estilo_base["cursiva"])
                        run.font.color.rgb = RGBColor(*color_texto)

                    if estilo_base["alineacion"]:
                        aplicar_alineacion_parrafo(parrafo, estilo_base["alineacion"])

def aplicar_estilo_a_texto(doc, texto_objetivo, estilo):
    if not texto_objetivo:
        return

    fuente = estilo["tipografia"]
    size_letra = int(estilo["tamanio_letra"])
    color_texto = convertir_color_bd_a_rgb(estilo["color_letra"])

    for parrafo in doc.paragraphs:
        if texto_objetivo in parrafo.text:
            for run in parrafo.runs:
                run.font.name = fuente
                run.font.size = Pt(size_letra)
                run.font.bold = bool(estilo["negrita"])
                run.font.italic = bool(estilo["cursiva"])
                run.font.color.rgb = RGBColor(*color_texto)

                if estilo["color_fondo"]:
                    colorear_fondo_run(run, estilo["color_fondo"])

            if estilo["alineacion"]:
                aplicar_alineacion_parrafo(parrafo, estilo["alineacion"])

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    if texto_objetivo in parrafo.text:
                        for run in parrafo.runs:
                            run.font.name = fuente
                            run.font.size = Pt(size_letra)
                            run.font.bold = bool(estilo["negrita"])
                            run.font.italic = bool(estilo["cursiva"])
                            run.font.color.rgb = RGBColor(*color_texto)

                            if estilo["color_fondo"]:
                                colorear_fondo_run(run, estilo["color_fondo"])

                        if estilo["alineacion"]:
                            aplicar_alineacion_parrafo(parrafo, estilo["alineacion"])


def obtener_ruta_documento_word(ruta_plantilla_word, nombre_archivo):
    ruta_documento = Path(ruta_plantilla_word) / nombre_archivo

    if not ruta_documento.exists():
        raise FileNotFoundError(
            f"No existe la plantilla Word en la ruta: {ruta_documento}"
        )

    return ruta_documento


def crear_word_cotizacion(
    empresa_nombre,
    ruta_membrete_word,
    ruta_plantilla_word,
    nombre_cliente,
    nombre_programa,
    registros,
    fuente,
    size_letra,
    paleta,
    estilos_bd
):
    ruta_plantilla = Path(ruta_plantilla_word) / "3.COTIZACION-FINAL.docx"

    if not ruta_plantilla.exists():
        raise FileNotFoundError(
            f"No existe la plantilla Word en la ruta: {ruta_plantilla}"
        )

    doc = crear_doc_con_membrete_y_contenido(
        ruta_membrete_word,
        ruta_plantilla
    )

    fecha_mas_antigua = obtener_fecha_mas_antigua_registros(registros)

    fecha_documento = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=4
    )

    fecha_str = f"Mérida, Yucatán, México a {fecha_documento}"

    tabla, total_general = crear_tabla_cotizacion(
        doc,
        registros,
        fuente,
        size_letra,
        paleta
    )

    parrafo_tabla = buscar_parrafo_con_texto(doc, "{{TABLA_COTIZACION}}")

    if parrafo_tabla is None:
        raise ValueError(
            "La plantilla Word no tiene el parámetro {{TABLA_COTIZACION}}."
        )

    insertar_tabla_despues_de_parrafo(parrafo_tabla, tabla)
    eliminar_parrafo(parrafo_tabla)
    titulo_1 = "COTIZACIÓN FINAL"

    reemplazos = {
        "{{FECHA}}": fecha_str,
        "{{TITULO_1}}": titulo_1,
        "{{EMPRESA_RECIBE}}": nombre_cliente,
        "{{EMPRESA_BRINDA}}": empresa_nombre,
        "{{NOMBRE_PROGRAMA}}": nombre_programa,
        "{{TOTAL}}": formatear_moneda(total_general),
        "{{TOTAL_LETRA}}": convertir_numero_a_letras_mxn(total_general),
        "{{FIRMA_EMPRESA}}": empresa_nombre,
    }

    reemplazar_parametros_documento(doc, reemplazos)

    if "texto_normal" in estilos_bd:
        aplicar_estilo_base_a_todo(doc, estilos_bd["texto_normal"])
    else:
        aplicar_fuente_base_a_todo(doc, fuente, size_letra, paleta["texto"])

    estilo_base = estilos_bd.get("texto_normal")

    aplicar_negrita_a_texto(
        doc,
        reemplazos.get("{{EMPRESA_RECIBE}}"),
        estilo_base
    )

    aplicar_negrita_a_texto(
        doc,
        reemplazos.get("{{EMPRESA_BRINDA}}"),
        estilo_base
    )

    if "titulo_1" in estilos_bd:
        aplicar_estilo_a_texto(doc, titulo_1, estilos_bd["titulo_1"])

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def crear_word_cotizacion_inicial(
    empresa_nombre,
    ruta_membrete_word,
    ruta_plantilla_word,
    nombre_cliente,
    nombre_programa,
    registros,
    fuente,
    size_letra,
    paleta,
    estilos_bd
):
    ruta_plantilla = Path(ruta_plantilla_word) / ARCHIVO_COTIZACION_INICIAL

    if not ruta_plantilla.exists():
        raise FileNotFoundError(
            f"No existe la plantilla Word en la ruta: {ruta_plantilla}"
        )

    doc = crear_doc_con_membrete_y_contenido(
        ruta_membrete_word,
        ruta_plantilla
    )

    registros_iniciales = crear_registros_cotizacion_inicial(registros)

    fecha_mas_antigua = obtener_fecha_mas_antigua_registros(registros)

    fecha_documento = calcular_fecha_larga_habil_mexico(
        fecha_str=fecha_mas_antigua,
        dias_restar=15,
        dias_sumar=2
    )

    fecha_str = f"Mérida, Yucatán, México a {fecha_documento}"

    tabla, total_general = crear_tabla_cotizacion(
        doc,
        registros_iniciales,
        fuente,
        size_letra,
        paleta
    )

    parrafo_tabla = buscar_parrafo_con_texto(doc, "{{TABLA_COTIZACION}}")

    if parrafo_tabla is None:
        raise ValueError(
            "La plantilla Word no tiene el parámetro {{TABLA_COTIZACION}}."
        )

    insertar_tabla_despues_de_parrafo(parrafo_tabla, tabla)
    eliminar_parrafo(parrafo_tabla)

    titulo_1 = "COTIZACIÓN INICIAL"

    reemplazos = {
        "{{FECHA}}": fecha_str,
        "{{TITULO_1}}": titulo_1,
        "{{EMPRESA_RECIBE}}": nombre_cliente,
        "{{EMPRESA_BRINDA}}": empresa_nombre,
        "{{NOMBRE_PROGRAMA}}": nombre_programa,
        "{{TOTAL}}": formatear_moneda(total_general),
        "{{TOTAL_LETRA}}": convertir_numero_a_letras_mxn(total_general),
        "{{FIRMA_EMPRESA}}": empresa_nombre,
    }

    reemplazar_parametros_documento(doc, reemplazos)

    if "texto_normal" in estilos_bd:
        aplicar_estilo_base_a_todo(doc, estilos_bd["texto_normal"])
    else:
        aplicar_fuente_base_a_todo(doc, fuente, size_letra, paleta["texto"])

    estilo_base = estilos_bd.get("texto_normal")

    aplicar_negrita_a_texto(doc, reemplazos.get("{{EMPRESA_RECIBE}}"), estilo_base)
    aplicar_negrita_a_texto(doc, reemplazos.get("{{EMPRESA_BRINDA}}"), estilo_base)
    aplicar_negrita_a_texto(doc, reemplazos.get("{{NOMBRE_PROGRAMA}}"), estilo_base)

    if "titulo_1" in estilos_bd:
        aplicar_estilo_a_texto(doc, titulo_1, estilos_bd["titulo_1"])

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def aplicar_fuente_base_a_todo(doc, fuente, size_letra, color_texto):
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            run.font.name = fuente
            run.font.size = Pt(size_letra)
            run.font.color.rgb = RGBColor(*color_texto)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.name = fuente
                        run.font.size = Pt(size_letra)
                        run.font.color.rgb = RGBColor(*color_texto)


def aplicar_alineacion_parrafo(parrafo, alineacion):
    alineacion = alineacion.lower().strip()

    if alineacion in ["izquierda", "left"]:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alineacion in ["centro", "center", "centrado"]:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alineacion in ["derecha", "right"]:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alineacion in ["justificado", "justify"]:
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def colorear_fondo_run(run, color_hex):
    color_hex = color_hex.replace("#", "")

    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    rPr.append(shd)


MESES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


def fecha_larga_es(fecha):
    return f"{fecha.day} de {MESES_ES[fecha.month]} del {fecha.year}"


def obtener_ruta_membrete(ruta_plantilla_word):
    ruta = Path(ruta_plantilla_word)

    if ruta.is_file():
        return ruta

    posibles = list(ruta.glob("plantilla*.docx"))

    if posibles:
        return posibles[0]

    raise FileNotFoundError(
        f"No se encontró una plantilla tipo 'plantilla*.docx' en: {ruta}"
    )


def limpiar_cuerpo_documento(doc):
    body = doc._body._element

    for elemento in list(body):
        if elemento.tag.endswith("sectPr"):
            continue

        body.remove(elemento)


def insertar_contenido(destino, origen):
    body_destino = destino._body._element
    body_origen = origen._body._element
    sectPr = body_destino.sectPr

    for elemento in body_origen:
        if elemento.tag.endswith("sectPr"):
            continue

        nuevo_elemento = deepcopy(elemento)

        if sectPr is not None:
            body_destino.insert(body_destino.index(sectPr), nuevo_elemento)
        else:
            body_destino.append(nuevo_elemento)

    return destino


def crear_doc_con_membrete_y_contenido(ruta_membrete_word, ruta_contenido_word):
    ruta_membrete_word = Path(ruta_membrete_word)
    ruta_contenido_word = Path(ruta_contenido_word)

    if not ruta_membrete_word.exists():
        raise FileNotFoundError(
            f"No existe el membrete Word: {ruta_membrete_word.resolve()}"
        )

    if not ruta_contenido_word.exists():
        raise FileNotFoundError(
            f"No existe la plantilla de contenido Word: {ruta_contenido_word.resolve()}"
        )

    doc_membrete = Document(str(ruta_membrete_word))
    doc_contenido = Document(str(ruta_contenido_word))

    body_membrete = doc_membrete._body._element
    sectPr_membrete = deepcopy(body_membrete.sectPr)

    limpiar_cuerpo_documento(doc_membrete)

    if body_membrete.sectPr is None and sectPr_membrete is not None:
        body_membrete.append(sectPr_membrete)

    insertar_contenido(doc_membrete, doc_contenido)

    return doc_membrete


def procesar_textboxes_xml(doc, reemplazos):
    for t in doc._element.xpath(".//*[local-name()='t']"):
        if t.text:
            for clave, valor in reemplazos.items():
                t.text = t.text.replace(clave, str(valor))


def obtener_periodo_servicio(registros):
    fechas = []

    for r in registros:
        fecha = pd.to_datetime(r.get("Fecha"), errors="coerce", dayfirst=True)
        if pd.notna(fecha):
            fechas.append(fecha.date())

    if not fechas:
        return ""

    fecha_inicio = min(fechas)
    fecha_fin = max(fechas)

    return f"{fecha_larga_es(fecha_inicio)} al {fecha_larga_es(fecha_fin)}"


def obtener_conceptos_unicos(registros):
    conceptos = []

    for r in registros:
        concepto = limpiar_texto(r.get("Concepto"))

        if concepto and concepto not in conceptos:
            conceptos.append(concepto)

    return conceptos


def reemplazar_conceptos_documento(doc, conceptos):
    texto_conceptos = "\n".join(
        f"{i + 1}.     {concepto}"
        for i, concepto in enumerate(conceptos)
    )

    reemplazos = {
        "{{CONCEPTOS_DESARROLLADOS}}": texto_conceptos,
    }

    reemplazar_parametros_documento(doc, reemplazos)


def formatear_fecha_corta(fecha):
    return f"{fecha.day}/{fecha.month}/{fecha.year}"


def crear_registros_calendario(registros):
    bloques = []
    bloque_actual = None

    for r in registros:
        concepto = limpiar_texto(r.get("Concepto"))
        fecha = pd.to_datetime(
            r.get("Fecha"),
            errors="coerce",
            dayfirst=True
        )

        if not concepto or pd.isna(fecha):
            continue

        fecha = fecha.date()

        if bloque_actual is None:
            bloque_actual = {
                "concepto": concepto,
                "fechas": [fecha],
            }
            bloques.append(bloque_actual)

        elif concepto == bloque_actual["concepto"]:
            bloque_actual["fechas"].append(fecha)

        else:
            bloque_actual = {
                "concepto": concepto,
                "fechas": [fecha],
            }
            bloques.append(bloque_actual)

    for i, bloque in enumerate(bloques):
        fecha_inicio = min(bloque["fechas"])

        if i < len(bloques) - 1:
            siguiente_inicio = min(bloques[i + 1]["fechas"])

            if siguiente_inicio == fecha_inicio:
                fecha_fin = fecha_inicio
            else:
                fecha_fin = siguiente_inicio - timedelta(days=1)
        else:
            fecha_fin = max(bloque["fechas"])

        bloque["fecha_inicio"] = fecha_inicio
        bloque["fecha_fin"] = fecha_fin

    return bloques


def fijar_ancho_celda(celda, ancho_pulgadas):
    celda.width = Inches(ancho_pulgadas)

    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()

    tcW = tcPr.first_child_found_in("w:tcW")

    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)

    tcW.set(qn("w:w"), str(int(ancho_pulgadas * 1440)))
    tcW.set(qn("w:type"), "dxa")


def crear_tabla_calendario(doc, registros_calendario, fuente, size_letra, paleta):
    tabla = doc.add_table(rows=1, cols=4)
    poner_bordes_tabla(tabla)

    tabla.autofit = False

    anchos_columnas = [0.35, 1.05, 1.45, 4.25]

    encabezados = [
        "#",
        "INICIO",
        "FINALIZACIÓN",
        "ACTIVIDAD A REALIZAR"
    ]

    header = tabla.rows[0]

    for i, texto in enumerate(encabezados):
        configurar_texto_celda(
            header.cells[i],
            texto,
            fuente,
            size_letra,
            True,
            paleta["texto"],
            WD_ALIGN_PARAGRAPH.LEFT,
        )
        colorear_celda(header.cells[i], paleta["principal"])
        fijar_ancho_celda(header.cells[i], anchos_columnas[i])

    for i, item in enumerate(registros_calendario, start=1):
        row = tabla.add_row()

        configurar_texto_celda(
            row.cells[0],
            str(i),
            fuente,
            size_letra,
            False,
            paleta["texto"],
            WD_ALIGN_PARAGRAPH.CENTER,
        )

        configurar_texto_celda(
            row.cells[1],
            formatear_fecha_corta(item["fecha_inicio"]),
            fuente,
            size_letra,
            False,
            paleta["texto"],
            WD_ALIGN_PARAGRAPH.LEFT,
        )

        configurar_texto_celda(
            row.cells[2],
            formatear_fecha_corta(item["fecha_fin"]),
            fuente,
            size_letra,
            False,
            paleta["texto"],
            WD_ALIGN_PARAGRAPH.LEFT,
        )

        configurar_texto_celda(
            row.cells[3],
            item["concepto"].upper(),
            fuente,
            size_letra,
            False,
            paleta["texto"],
            WD_ALIGN_PARAGRAPH.LEFT,
        )

        for idx_col, ancho in enumerate(anchos_columnas):
            fijar_ancho_celda(row.cells[idx_col], ancho)

    return tabla


def insertar_tabla_calendario_en_doc(doc, registros_calendario, fuente, size_letra, paleta):
    parrafo_tabla = buscar_parrafo_con_texto(
        doc,
        "{{TABLA_CALENDARIO}}"
    )

    if parrafo_tabla is None:
        return

    tabla = crear_tabla_calendario(
        doc,
        registros_calendario,
        fuente,
        size_letra,
        paleta
    )

    insertar_tabla_despues_de_parrafo(parrafo_tabla, tabla)
    eliminar_parrafo(parrafo_tabla)


def insertar_parrafo_despues(parrafo, texto="", estilo=None):
    nuevo_elemento = OxmlElement("w:p")
    parrafo._p.addnext(nuevo_elemento)

    nuevo_parrafo = Paragraph(nuevo_elemento, parrafo._parent)

    if estilo:
        nuevo_parrafo.style = estilo

    if texto:
        nuevo_parrafo.add_run(texto)

    return nuevo_parrafo


def insertar_desarrollo_conceptos_entregable(doc, conceptos, textos_por_concepto, estilos_bd):
    parrafo_marker = buscar_parrafo_con_texto(
        doc,
        "{{DESARROLLO_CONCEPTOS}}"
    )

    if parrafo_marker is None:
        raise ValueError(
            "La plantilla Word no tiene el parámetro {{DESARROLLO_CONCEPTOS}}."
        )

    ultimo_parrafo = parrafo_marker

    for concepto in conceptos:
        titulo = concepto.upper()

        parrafo_titulo = insertar_parrafo_despues(
            ultimo_parrafo,
            titulo
        )

        aplicar_heading_a_texto(
            doc,
            titulo,
            nivel=1
        )

        if "titulo_1" in estilos_bd:
            aplicar_estilo_a_texto(
                doc,
                titulo,
                estilos_bd["titulo_1"]
            )
        crear_word_propuesta_desde_plantilla
        texto = textos_por_concepto.get(concepto, "")

        parrafo_texto = insertar_parrafo_despues(
            parrafo_titulo,
            texto
        )

        ultimo_parrafo = parrafo_texto

    eliminar_parrafo(parrafo_marker)


def crear_documento_desde_plantilla(
    ruta_membrete_word,
    ruta_plantilla_word,
    nombre_archivo_contenido,
    reemplazos,
    estilos_bd,
    conceptos=None,
    nombre_cliente=None,
    empresa_nombre=None,
    registros_calendario=None,
    fuente=None,
    size_letra=None,
    paleta=None,
):
    ruta_documento = (
        Path(ruta_plantilla_word) / nombre_archivo_contenido
    )

    if not ruta_documento.exists():
        raise FileNotFoundError(
            f"No existe el documento Word3: {ruta_documento}"
        )

    doc = crear_doc_con_membrete_y_contenido(
        ruta_membrete_word,
        ruta_documento
    )

    reemplazar_parametros_documento(doc, reemplazos)

    procesar_textboxes_xml(doc, reemplazos)

    if conceptos:
        reemplazar_conceptos_documento(doc, conceptos)

    if registros_calendario:
        insertar_tabla_calendario_en_doc(
            doc,
            registros_calendario,
            fuente,
            size_letra,
            paleta
        )

    if "texto_normal" in estilos_bd:
        aplicar_estilo_base_a_todo(
            doc,
            estilos_bd["texto_normal"]
        )

    estilo_base = estilos_bd.get("texto_normal")

    aplicar_negrita_a_texto(
        doc,
        reemplazos.get("{{EMPRESA_RECIBE}}"),
        estilo_base
    )

    aplicar_negrita_a_texto(
        doc,
        reemplazos.get("{{EMPRESA_BRINDA}}"),
        estilo_base
    )

    aplicar_negrita_a_texto(
        doc,
        reemplazos.get("{{NOMBRE_PROGRAMA}}"),
        estilo_base
    )

    titulo_1 = reemplazos.get("{{TITULO_1}}")

    if titulo_1 and "titulo_1" in estilos_bd:
        if titulo_1:
            aplicar_heading_a_texto(
                doc,
                titulo_1,
                nivel=1
            )

        aplicar_estilo_a_texto(
            doc,
            titulo_1,
            estilos_bd["titulo_1"]
        )

    output = BytesIO()

    doc.save(output)

    output.seek(0)

    return output


def crear_word_propuesta_desde_plantilla(
    ruta_membrete_word,
    ruta_plantilla_word,
    reemplazos,
    estilos_bd,
    fuente,
    size_letra,
    paleta,
):
    ruta_documento = Path(ruta_plantilla_word) / ARCHIVO_PROPUESTA

    if not ruta_documento.exists():
        raise FileNotFoundError(
            f"No existe la plantilla de propuesta: {ruta_documento}"
        )

    doc = crear_doc_con_membrete_y_contenido(
        ruta_membrete_word,
        ruta_documento
    )

    reemplazar_parametros_documento(doc, reemplazos)
    procesar_textboxes_xml(doc, reemplazos)

    if "texto_normal" in estilos_bd:
        aplicar_estilo_base_a_todo(doc, estilos_bd["texto_normal"])
    else:
        aplicar_fuente_base_a_todo(doc, fuente, size_letra, paleta["texto"])

    estilo_base = estilos_bd.get("texto_normal")

    aplicar_negrita_a_texto(doc, reemplazos.get("{{EMPRESA_RECIBE}}"), estilo_base)
    aplicar_negrita_a_texto(doc, reemplazos.get("{{EMPRESA_BRINDA}}"), estilo_base)
    aplicar_negrita_a_texto(doc, reemplazos.get("{{NOMBRE_PROGRAMA}}"), estilo_base)

    titulo_1 = reemplazos.get("{{TITULO_1}}")

    if titulo_1:
        aplicar_heading_a_texto(doc, titulo_1, nivel=1)

    if titulo_1 and "titulo_1" in estilos_bd:
        aplicar_estilo_a_texto(doc, titulo_1, estilos_bd["titulo_1"])

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def crear_word_entregable(
    ruta_membrete_word,
    ruta_plantilla_word,
    reemplazos,
    estilos_bd,
    conceptos,
    textos_por_concepto,
    registros,
    fuente,
    size_letra,
    paleta,
):
    ruta_documento = Path(ruta_plantilla_word) / ARCHIVO_ENTREGABLE

    if not ruta_documento.exists():
        raise FileNotFoundError(
            f"No existe la plantilla de entregable: {ruta_documento}"
        )

    doc = crear_doc_con_membrete_y_contenido(
        ruta_membrete_word,
        ruta_documento
    )

    tabla_pagos, total_general = crear_tabla_cotizacion(
        doc,
        registros,
        fuente,
        size_letra,
        paleta
    )

    reemplazos["{{TOTAL}}"] = formatear_moneda(total_general)
    reemplazos["{{TOTAL_LETRA}}"] = convertir_numero_a_letras_mxn(total_general)

    reemplazar_parametros_documento(doc, reemplazos)
    procesar_textboxes_xml(doc, reemplazos)

    insertar_desarrollo_conceptos_entregable(
        doc,
        conceptos,
        textos_por_concepto,
        estilos_bd
    )

    parrafo_control = buscar_parrafo_con_texto(
        doc,
        "{{CONTROL_PAGOS}}"
    )

    if parrafo_control:
        insertar_tabla_despues_de_parrafo(
            parrafo_control,
            tabla_pagos
        )
        eliminar_parrafo(parrafo_control)

    if "texto_normal" in estilos_bd:
        aplicar_estilo_base_a_todo(
            doc,
            estilos_bd["texto_normal"]
        )
    else:
        aplicar_fuente_base_a_todo(
            doc,
            fuente,
            size_letra,
            paleta["texto"]
        )

    for clave in [
        "{{EMPRESA_RECIBE}}",
        "{{EMPRESA_BRINDA}}",
        "{{NOMBRE_PROGRAMA}}"
    ]:
        aplicar_negrita_a_texto(
            doc,
            reemplazos.get(clave),
            estilos_bd.get("texto_normal")
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def mostrar_modulo_cotizacion_final():
    st.title("Entregables AA")

    empresas = obtener_empresas_con_plantilla()

    if not empresas:
        st.error(
            "No hay empresas con plantilla y membrete asignados."
        )
        return

    nombre_programa = st.text_input(
        "Nombre del programa"
    )

    col1, col2 = st.columns(2)

    with col1:
        empresa_seleccionada = st.selectbox(
            "Empresa prestadora del servicio",
            list(empresas.keys()),
        )

    with col2:
        nombre_cliente = st.text_input(
            "Empresa que recibe"
        )

    col3, col4, col5 = st.columns(3)

    with col3:
        archivo = st.file_uploader(
            "Excel de conceptos",
            type=["xlsx", "xls"],
        )

    empresa_data = empresas[empresa_seleccionada]

    plantilla_id = empresa_data["plantilla_id"]

    fuente = empresa_data["tipografia_base"]
    size_letra = int(empresa_data["tamanio_base"])

    color_texto_base = convertir_color_bd_a_rgb(
        empresa_data["color_texto_base"]
    )

    color_primario = convertir_color_bd_a_rgb(
        empresa_data["color_primario"]
    )

    color_secundario = convertir_color_bd_a_rgb(
        empresa_data["color_secundario"]
    )

    paleta = {
        "principal": color_primario,
        "secundario": color_secundario,
        "texto": color_texto_base,
    }

    estilos_bd = obtener_estilos_word_por_plantilla(plantilla_id)

    nombre_empresa = empresa_data["razon_social"]

    ruta_membrete_word = empresa_data["membrete_path"]
    ruta_plantilla_word = empresa_data["plantilla_path"]

    color_primario = convertir_color_bd_a_rgb(empresa_data["color_primario"])
    color_secundario = convertir_color_bd_a_rgb(empresa_data["color_secundario"])

    paleta = {
        "principal": color_primario,
        "secundario": color_secundario,
        "texto": (0, 0, 0),
    }

    if archivo is None:
        st.warning("Sube un archivo Excel para continuar.")
        return

    if not nombre_cliente:
        st.warning("Escribe el nombre del cliente.")
        return

    if not nombre_programa:
        st.warning("Escribe el nombre del programa.")
        return

    try:
        registros = extraer_datos_cotizacion_excel(archivo)

        archivo.seek(0)

        registros_calendario = extraer_datos_calendario_excel(archivo)

        fechas_entregables = calcular_fechas_entregables(registros)

        word = crear_word_cotizacion(
            empresa_nombre=nombre_empresa,
            ruta_membrete_word=ruta_membrete_word,
            ruta_plantilla_word=ruta_plantilla_word,
            nombre_cliente=nombre_cliente,
            nombre_programa=nombre_programa,
            registros=registros,
            fuente=fuente,
            size_letra=size_letra,
            paleta=paleta,
            estilos_bd=estilos_bd,
        )

        word_inicial = crear_word_cotizacion_inicial(
            empresa_nombre=nombre_empresa,
            ruta_membrete_word=ruta_membrete_word,
            ruta_plantilla_word=ruta_plantilla_word,
            nombre_cliente=nombre_cliente,
            nombre_programa=nombre_programa,
            registros=registros,
            fuente=fuente,
            size_letra=size_letra,
            paleta=paleta,
            estilos_bd=estilos_bd,
        )

        conceptos_unicos = obtener_conceptos_unicos(registros)

        conceptos_factura = "\n".join(
            [f"{i + 1}. {c}" for i, c in enumerate(conceptos_unicos)]
        )

        lugar_fecha = fechas_entregables["propuesta"]

        periodo_servicio = obtener_periodo_servicio(registros)
        conceptos_unicos = obtener_conceptos_unicos(registros)

        reemplazos_acuse = {
            "{{TITULO_1}}": "ACUSE DE RECIBO",
            "{{EMPRESA_RECIBE}}": nombre_cliente.upper(),
            "{{EMPRESA_BRINDA}}": nombre_empresa.upper(),
            "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
            "{{FECHA_SERVICIO}}": periodo_servicio,
            "29 de septiembre del 2025 al 17 de marzo del 2026": periodo_servicio,
        }

        reemplazos_resumen = {
            "{{TITULO_1}}": "RESUMEN EJECUTIVO",
            "{{EMPRESA_RECIBE}}": nombre_cliente.upper(),
            "{{EMPRESA_BRINDA}}": nombre_empresa.upper(),
            "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
            "{{FECHA_SERVICIO}}": periodo_servicio,
            "29 de septiembre del 2025 al 17 de marzo del 2026": periodo_servicio,
        }

        fecha_calendario = fechas_entregables["calendario"]

        reemplazos_calendario = {
            "{{FECHA}}": fecha_calendario,
            "{{TITULO_1}}": "CALENDARIO DE TRABAJO",
            "{{EMPRESA_RECIBE}}": nombre_cliente.upper(),
            "{{EMPRESA_BRINDA}}": nombre_empresa.upper(),
            "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
            "{{FECHA_SERVICIO}}": periodo_servicio,
            "29 de septiembre del 2025 al 17 de marzo del 2026": periodo_servicio,
        }

        st.subheader("Descargas de documentos")

        col_doc1, col_doc2, col_doc3, col_doc4, col_doc5, col_doc6, col_doc7 = st.columns(7)

        with col_doc1:
            if st.button(
                    "Generar propuesta IA",
                    key="btn_generar_propuesta_ia",
                    use_container_width=True
            ):
                with st.spinner("Generando propuesta IA..."):
                    texto_propuesta = generar_objetivo_general(nombre_programa)

                    introduccion = generar_introduccion(
                        nombre_programa,
                        conceptos_factura
                    )

                    problematica = generar_problematica(
                        nombre_programa,
                        conceptos_factura
                    )

                    objetivo_general = generar_objetivo_general_tabla(
                        nombre_programa,
                        conceptos_factura,
                        nombre_cliente
                    )

                    objetivos_especificos = generar_objetivos_especificos(
                        nombre_programa,
                        conceptos_factura,
                        nombre_cliente
                    )

                    metodologia = generar_metodologia(
                        nombre_programa,
                        conceptos_factura
                    )

                    reemplazos_propuesta = {
                        "{{FECHA}}": lugar_fecha,
                        "{{TITULO_1}}": "PROPUESTA DE SERVICIOS",
                        "{{EMPRESA_RECIBE}}": nombre_cliente.upper(),
                        "{{EMPRESA_BRINDA}}": nombre_empresa.upper(),
                        "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
                        "{{TEXTO_PROPUESTA}}": texto_propuesta,
                        "{{INTRODUCCION}}": introduccion,
                        "{{PROBLEMATICA}}": problematica,
                        "{{OBJETIVO_GENERAL}}": objetivo_general,
                        "{{OBJETIVOS_ESPECIFICOS}}": objetivos_especificos,
                        "{{METODOLOGIA}}": metodologia,
                        "{{FIRMA_EMPRESA}}": nombre_empresa.upper(),
                    }

                    word_propuesta = crear_word_propuesta_desde_plantilla(
                        ruta_membrete_word=ruta_membrete_word,
                        ruta_plantilla_word=ruta_plantilla_word,
                        reemplazos=reemplazos_propuesta,
                        estilos_bd=estilos_bd,
                        fuente=fuente,
                        size_letra=size_letra,
                        paleta=paleta,
                    )

                    st.session_state["word_propuesta_ia"] = word_propuesta

            if "word_propuesta_ia" in st.session_state:
                st.download_button(
                    label="Descargar propuesta IA",
                    data=st.session_state["word_propuesta_ia"],
                    file_name="propuesta_ia.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_descargar_propuesta_ia",
                    on_click="ignore"
                )

        with col_doc2:
            st.download_button(
                label="Descargar cotización inicial",
                data=word_inicial,
                file_name="cotizacion_inicial.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_descargar_cotizacion_inicial",
                on_click = "ignore"
            )

        with col_doc3:
            st.download_button(
                label="Descargar cotización final",
                data=word,
                file_name="cotizacion_final.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_descargar_cotizacion_final",
                on_click="ignore"
            )

        with col_doc4:
            calendario = crear_documento_desde_plantilla(
                ruta_membrete_word=ruta_membrete_word,
                ruta_plantilla_word=ruta_plantilla_word,
                nombre_archivo_contenido=ARCHIVO_CALENDARIO,
                reemplazos=reemplazos_calendario,
                estilos_bd=estilos_bd,
                nombre_cliente=nombre_cliente.upper(),
                empresa_nombre=nombre_empresa.upper(),
                registros_calendario=registros_calendario,
                fuente=fuente,
                size_letra=size_letra,
                paleta=paleta,
            )

            st.download_button(
                label="Descargar calendario",
                data=calendario,
                file_name="calendario_de_trabajo.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_descargar_calendario",
                on_click="ignore"
            )

        with col_doc5:
            resumen = crear_documento_desde_plantilla(
                ruta_membrete_word=ruta_membrete_word,
                ruta_plantilla_word=ruta_plantilla_word,
                nombre_archivo_contenido=ARCHIVO_RESUMEN,
                reemplazos=reemplazos_resumen,
                estilos_bd=estilos_bd,
                conceptos=conceptos_unicos,
                nombre_cliente=nombre_cliente.upper(),
                empresa_nombre=nombre_empresa.upper(),
            )

            st.download_button(
                label="Descargar resumen",
                data=resumen,
                file_name="resumen_ejecutivo.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_descargar_resumen",
                on_click="ignore"
            )

        with col_doc6:
            acuse = crear_documento_desde_plantilla(
                ruta_membrete_word=ruta_membrete_word,
                ruta_plantilla_word=ruta_plantilla_word,
                nombre_archivo_contenido=ARCHIVO_ACUSE,
                reemplazos=reemplazos_acuse,
                estilos_bd=estilos_bd,
                nombre_cliente=nombre_cliente.upper(),
                empresa_nombre=nombre_empresa.upper(),
            )

            st.download_button(
                label="Descargar acuse",
                data=acuse,
                file_name="acuse.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_descargar_acuse",
                on_click="ignore"
            )

        with col_doc7:
            if st.button(
                    "Generar entregable",
                    key="btn_generar_entregable",
                    use_container_width=True
            ):
                with st.spinner("Generando entregable..."):
                    introduccion_entregable = generar_introduccion_entregable(
                        nombre_programa,
                        nombre_cliente,
                        nombre_empresa
                    )

                    objetivo_general_entregable = generar_objetivo_general_tabla(
                        nombre_programa,
                        conceptos_factura,
                        nombre_cliente
                    )

                    objetivos_especificos_entregable = generar_objetivos_especificos(
                        nombre_programa,
                        conceptos_factura,
                        nombre_cliente
                    )

                    finalizacion_servicio = (
                        f"{nombre_cliente.upper()}, doy por terminada la prestación de servicio "
                        f"profesional mediante el cual obtuve un documento que proporciona una "
                        f"gestión documental más organizada, segura y efectiva, adaptada a las "
                        f"necesidades actuales de transformación digital. Dicho documento me fue "
                        f"proporcionado por {nombre_empresa.upper()} en Mérida, Yucatán, México."
                    )

                    referencias = generar_referencias_entregable(
                        nombre_programa,
                        conceptos_unicos
                    )

                    textos_por_concepto = {}

                    for concepto in conceptos_unicos:
                        textos_por_concepto[concepto] = generar_texto_concepto_entregable(
                            nombre_programa,
                            concepto
                        )

                    texto_servicio_prestado = (
                        f"El servicio que se presenta fue ofrecido en las instalaciones de "
                        f"{nombre_cliente.upper()} por la empresa {nombre_empresa.upper()}.\n\n"
                        f"El personal encargado de brindar el servicio “{nombre_programa.upper()}”, "
                        f"tiene la experiencia, capacidad y pericia para poder plasmar sus conocimientos "
                        f"en este documento.\n\n"
                        f"El servicio se llevó a cabo en las instalaciones del cliente, durante el periodo "
                        f"comprendido del {periodo_servicio}. Durante su desarrollo, se llevaron a cabo "
                        f"las siguientes actividades:"
                    )

                    reemplazos_entregable = {
                        "{{FECHA}}": fechas_entregables["cotizacion_final"],
                        "{{EMPRESA_RECIBE}}": nombre_cliente.upper(),
                        "{{EMPRESA_BRINDA}}": nombre_empresa.upper(),
                        "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
                        "{{PERIODO_SERVICIO}}": periodo_servicio,
                        "{{SERVICIO_PRESTADO}}": texto_servicio_prestado,
                        "{{INTRODUCCION}}": introduccion_entregable,
                        "{{OBJETIVO_GENERAL}}": objetivo_general_entregable,
                        "{{OBJETIVOS_ESPECIFICOS}}": objetivos_especificos_entregable,
                        "{{FINALIZACION_SERVICIO}}": finalizacion_servicio,
                        "{{REFERENCIAS}}": referencias,
                    }

                    word_entregable = crear_word_entregable(
                        ruta_membrete_word=ruta_membrete_word,
                        ruta_plantilla_word=ruta_plantilla_word,
                        reemplazos=reemplazos_entregable,
                        estilos_bd=estilos_bd,
                        conceptos=conceptos_unicos,
                        textos_por_concepto=textos_por_concepto,
                        registros=registros,
                        fuente=fuente,
                        size_letra=size_letra,
                        paleta=paleta,
                    )

                    st.session_state["word_entregable"] = word_entregable

            if "word_entregable" in st.session_state:
                st.download_button(
                    label="Descargar entregable",
                    data=st.session_state["word_entregable"],
                    file_name="entregable.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_descargar_entregable",
                    on_click="ignore"
                )

    except ValueError as e:
        st.error(str(e))

    except FileNotFoundError as e:
        st.error(str(e))

    except Exception as e:
        st.error(f"Ocurrió un error: {e}")


