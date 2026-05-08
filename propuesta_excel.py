import re
from io import BytesIO
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from datetime import timedelta, date


def limpiar_texto(valor):
    if pd.isna(valor):
        return ""
    return str(valor).strip()


def extraer_entre_parentesis(texto):
    texto = limpiar_texto(texto)
    encontrados = re.findall(r"\((.*?)\)", texto)
    if encontrados:
        return encontrados[-1].strip()
    return texto


def limpiar_objetivo(texto):
    texto = limpiar_texto(texto)
    return re.sub(r"^Objetivo:\s*", "", texto, flags=re.IGNORECASE).strip()


def formatear_fecha(valor):
    if pd.isna(valor) or valor == "":
        return ""

    if isinstance(valor, datetime):
        return valor.strftime("%d/%m/%Y")

    try:
        fecha = pd.to_datetime(valor, errors="coerce", dayfirst=True)
        if pd.notna(fecha):
            return fecha.strftime("%d/%m/%Y")
    except Exception:
        pass

    return str(valor)


def extraer_datos_excel(archivo_excel):
    df = pd.read_excel(archivo_excel, header=None)

    cliente = ""
    programa = ""
    objetivo = ""

    for i in range(len(df)):
        fila = df.iloc[i].astype(str).tolist()

        for j, valor in enumerate(fila):
            valor_limpio = limpiar_texto(valor)

            if valor_limpio.lower() == "cliente:":
                cliente = limpiar_texto(df.iloc[i, j + 1])

            elif valor_limpio.lower() == "nombre del programa:":
                programa = limpiar_texto(df.iloc[i, j + 1])

            elif valor_limpio.lower().startswith("objetivo:"):
                objetivo = limpiar_objetivo(valor_limpio)

    fila_encabezados = None

    for i in range(len(df)):
        fila = [limpiar_texto(x).lower() for x in df.iloc[i].tolist()]
        if "concepto" in fila and "total factura ($)" in fila and "fecha" in fila:
            fila_encabezados = i
            break

    if fila_encabezados is None:
        raise ValueError("No se encontró la fila de encabezados con Concepto, Total Factura ($) y Fecha.")

    encabezados = [limpiar_texto(x) for x in df.iloc[fila_encabezados].tolist()]

    col_concepto = encabezados.index("Concepto")
    col_monto = encabezados.index("Total Factura ($)")
    col_fecha = encabezados.index("Fecha")

    registros = []

    ultimo_concepto = ""

    for i in range(fila_encabezados + 1, len(df)):
        concepto_raw = limpiar_texto(df.iloc[i, col_concepto])
        monto = df.iloc[i, col_monto]
        fecha = df.iloc[i, col_fecha]

        if concepto_raw:
            ultimo_concepto = extraer_entre_parentesis(concepto_raw)

        if pd.notna(monto):
            registros.append({
                "Metodología / Concepto": ultimo_concepto,
                "Monto": float(monto),
                "Fecha": formatear_fecha(fecha)
            })

    return cliente, programa, objetivo, registros

from datetime import timedelta, date


def obtener_dias_festivos_mexico(anio):
    """
    Aquí ajustas los días inhábiles oficiales de México.
    Agrega o quita fechas según tu criterio.
    """

    festivos = set()

    # Fechas fijas
    festivos.add(date(anio, 1, 1))    # Año Nuevo
    festivos.add(date(anio, 5, 1))    # Día del Trabajo
    festivos.add(date(anio, 9, 16))   # Independencia
    festivos.add(date(anio, 12, 25))  # Navidad

    # Primer lunes de febrero - Constitución
    festivos.add(primer_lunes(anio, 2))

    # Tercer lunes de marzo - Natalicio de Benito Juárez
    festivos.add(tercer_lunes(anio, 3))

    # Tercer lunes de noviembre - Revolución Mexicana
    festivos.add(tercer_lunes(anio, 11))

    return festivos


def primer_lunes(anio, mes):
    d = date(anio, mes, 1)
    while d.weekday() != 0:
        d += timedelta(days=1)
    return d


def tercer_lunes(anio, mes):
    d = primer_lunes(anio, mes)
    return d + timedelta(days=14)


def es_dia_habil_mexico(fecha):
    """
    Se consideran inhábiles:
    - Domingos
    - Días festivos oficiales de México
    """

    if fecha.weekday() == 6:  # Domingo
        return False

    festivos = obtener_dias_festivos_mexico(fecha.year)
    return fecha not in festivos


def restar_15_dias_habiles_mexico(fecha_base):
    contador = 0
    fecha_actual = fecha_base

    while contador < 15:
        fecha_actual -= timedelta(days=1)

        if es_dia_habil_mexico(fecha_actual):
            contador += 1

    return fecha_actual


def convertir_fecha_a_date(valor):
    if isinstance(valor, date):
        return valor

    fecha = pd.to_datetime(valor, errors="coerce", dayfirst=True)

    if pd.isna(fecha):
        return None

    return fecha.date()


def obtener_fecha_mas_antigua_y_fecha_retroactiva(registros):
    fechas = []

    for r in registros:
        fecha = convertir_fecha_a_date(r["Fecha"])
        if fecha:
            fechas.append(fecha)

    if not fechas:
        return "", ""

    fecha_mas_antigua = min(fechas)
    fecha_15_habiles_antes = restar_15_dias_habiles_mexico(fecha_mas_antigua)

    return fecha_mas_antigua, fecha_15_habiles_antes
def generar_word(cliente, programa, objetivo, registros):
    doc = Document()

    fecha_mas_antigua, fecha_15_habiles_antes = obtener_fecha_mas_antigua_y_fecha_retroactiva(registros)

    doc.add_heading("DATOS EXTRAÍDOS DEL PROYECTO", level=1)

    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Nombre del programa: {programa}")
    doc.add_paragraph(f"Objetivo: {objetivo}")

    if fecha_mas_antigua:
        doc.add_paragraph(
            f"Primera fecha de factura: {fecha_mas_antigua.strftime('%d/%m/%Y')}"
        )

    if fecha_15_habiles_antes:
        doc.add_paragraph(
            f"Fecha considerando 15 días hábiles anteriores: {fecha_15_habiles_antes.strftime('%d/%m/%Y')}"
        )

    doc.add_heading("Conceptos de factura", level=2)

    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"

    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Metodología / Concepto"
    encabezados[1].text = "Monto"
    encabezados[2].text = "Fecha"

    for r in registros:
        fila = tabla.add_row().cells
        fila[0].text = r["Metodología / Concepto"]
        fila[1].text = f"${r['Monto']:,.2f}"
        fila[2].text = r["Fecha"]

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def mostrar_modulo_propuesta_excel():
    st.title("Extraer datos de propuesta desde Excel")

    archivo = st.file_uploader(
        "Sube el archivo Excel",
        type=["xlsx", "xls"]
    )

    if archivo is None:
        st.info("Sube un archivo Excel para extraer los datos.")
        return

    try:
        cliente, programa, objetivo, registros = extraer_datos_excel(archivo)

        st.subheader("Vista previa de datos generales")
        st.write(f"**Cliente:** {cliente}")
        st.write(f"**Nombre del programa:** {programa}")
        st.write(f"**Objetivo:** {objetivo}")

        st.subheader("Vista previa de conceptos, montos y fechas")

        df_preview = pd.DataFrame(registros)
        df_preview["Monto"] = df_preview["Monto"].map(lambda x: f"${x:,.2f}")

        st.dataframe(df_preview, use_container_width=True)

        word = generar_word(cliente, programa, objetivo, registros)

        st.download_button(
            label="Descargar Word",
            data=word,
            file_name="datos_extraidos_propuesta.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"Ocurrió un error al procesar el Excel: {e}")