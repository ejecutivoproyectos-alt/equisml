import os
import re
import tempfile
import unicodedata
from datetime import datetime, date, timedelta

import streamlit as st
from openpyxl import load_workbook, Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

STATUS_COLORS = {
    "LABORADO": "19A7D8",
    "SI LABORO": "19A7D8",
    "SI LABORÓ": "19A7D8",
    "NO LABORO": "C0504D",
    "NO LABORÓ": "C0504D",
    "INHABIL": "FF0000",
    "INHÁBIL": "FF0000",
    "INCAPACIDAD": "59A8C0",
    "VACACIONES": "8C8C8C",
    "BAJA": "FFFF00",
    "NO HABIA INGRESADO": "FFFFFF",
    "NO HABÍA INGRESADO": "FFFFFF"
}

STATUS_LABELS = [
    ("SI LABORÓ", "19A7D8"),
    ("NO LABORÓ", "C0504D"),
    ("INHÁBIL", "FF0000"),
    ("INCAPACIDAD", "59A8C0"),
    ("VACACIONES", "8C8C8C"),
    ("BAJA", "FFFF00"),
    ("NO HABÍA INGRESADO", "FFFFFF"),
]

DAY_LETTERS = {
    0: "L",
    1: "M",
    2: "M",
    3: "J",
    4: "V",
    5: "S",
    6: "D",
}

MONTH_NAMES_ES = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE",
}


def nth_weekday_of_month(year, month, weekday, n):
    first_day = date(year, month, 1)
    days_until_weekday = (weekday - first_day.weekday()) % 7
    first_occurrence = first_day + timedelta(days=days_until_weekday)
    return first_occurrence + timedelta(weeks=n - 1)


def get_official_mexico_holidays(year):
    holidays = set()

    holidays.add(date(year, 1, 1))
    holidays.add(nth_weekday_of_month(year, 2, 0, 1))
    holidays.add(nth_weekday_of_month(year, 3, 0, 3))
    holidays.add(date(year, 5, 1))
    holidays.add(date(year, 9, 16))
    holidays.add(nth_weekday_of_month(year, 11, 0, 3))

    if year >= 2024 and (year - 2024) % 6 == 0:
        holidays.add(date(year, 10, 1))

    holidays.add(date(year, 12, 25))

    return holidays


def normalize_text(text):
    if text is None:
        return ""
    text = str(text).strip().upper()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = " ".join(text.split())
    return text


def parse_excel_date(value):
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value

    if isinstance(value, str):
        value = value.strip()
        formats = [
            "%d-%m-%Y",
            "%d/%m/%Y",
            "%Y-%m-%d",
            "%d-%m-%y",
            "%d/%m/%y",
        ]
        for fmt in formats:
            try:
                return datetime.strptime(value, fmt).date()
            except ValueError:
                pass
    return None


def find_header_row_and_columns(ws):
    empleado_col = None
    cliente_col = None
    header_row = None
    date_cols = []

    for row in range(1, min(ws.max_row, 30) + 1):
        row_values = [ws.cell(row=row, column=col).value for col in range(1, ws.max_column + 1)]
        normalized = [normalize_text(v) for v in row_values]

        if "EMPLEADO" in normalized:
            header_row = row
            empleado_col = normalized.index("EMPLEADO") + 1
            if "CLIENTE" in normalized:
                cliente_col = normalized.index("CLIENTE") + 1

            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=header_row, column=col).value
                if parse_excel_date(cell_value):
                    date_cols.append(col)

            break

    if not header_row or not empleado_col or not date_cols:
        raise ValueError(
            "No se pudo identificar la estructura del archivo. "
            "Verifica que exista una fila con 'Empleado' y encabezados de fechas."
        )

    return header_row, empleado_col, cliente_col, date_cols


def split_employee_code_name(raw_name):
    text = "" if raw_name is None else str(raw_name).strip()
    text = " ".join(text.split())
    parts = re.split(r"\s*-\s*", text, maxsplit=1)

    if len(parts) == 2:
        code = parts[0].strip()
        name = parts[1].strip()
        return code, name

    return "", text


def build_employee_index(ws, header_row, empleado_col, cliente_col, date_cols):
    employees = {}

    for row in range(header_row + 1, ws.max_row + 1):
        raw_name = ws.cell(row=row, column=empleado_col).value
        if raw_name is None:
            continue

        original_name = " ".join(str(raw_name).strip().split())
        if not original_name:
            continue

        employee_code, clean_name = split_employee_code_name(original_name)

        norm_original = normalize_text(original_name)
        norm_clean = normalize_text(clean_name)
        norm_code = normalize_text(employee_code)

        cliente = ws.cell(row=row, column=cliente_col).value if cliente_col else ""

        statuses = {}
        for col in date_cols:
            fecha = parse_excel_date(ws.cell(row=header_row, column=col).value)
            estado = ws.cell(row=row, column=col).value
            if fecha:
                statuses[fecha] = normalize_text(estado)

        data = {
            "row": row,
            "original_name": original_name,
            "clean_name": clean_name,
            "employee_code": employee_code,
            "cliente": str(cliente).strip() if cliente else "",
            "statuses": statuses
        }

        if norm_original:
            employees[norm_original] = data
        if norm_clean:
            employees[norm_clean] = data
        if norm_code:
            employees[norm_code] = data

    return employees


def find_employee(employee_index, search_name):
    search = normalize_text(search_name)

    if not search:
        return None

    if search in employee_index:
        return employee_index[search]

    partial_matches = []
    seen_rows = set()

    for _, data in employee_index.items():
        row_id = data["row"]
        if row_id in seen_rows:
            continue

        seen_rows.add(row_id)

        original_norm = normalize_text(data.get("original_name", ""))
        clean_norm = normalize_text(data.get("clean_name", ""))
        code_norm = normalize_text(data.get("employee_code", ""))

        if (
            search in clean_norm
            or clean_norm in search
            or search in original_norm
            or search in code_norm
        ):
            partial_matches.append(data)

    if len(partial_matches) == 1:
        return partial_matches[0]

    return None


def get_months_present(date_list):
    return sorted(set((d.year, d.month) for d in date_list))


def get_pre_month_default_status(statuses, year, month):
    month_days = sorted(
        d for d in statuses.keys()
        if d.year == year and d.month == month and d.weekday() <= 5
    )

    for d in month_days:
        status = normalize_text(statuses.get(d, ""))
        if status:
            if status == "NO HABIA INGRESADO":
                return "NO HABÍA INGRESADO"
            return "SI LABORÓ"

    return "SI LABORÓ"


def status_to_fill(status):
    color = STATUS_COLORS.get(normalize_text(status), "FFFFFF")
    return PatternFill(fill_type="solid", fgColor=color)


def apply_border(cell):
    thin = Side(style="thin", color="000000")
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def set_cell(ws, row, col, value=None, fill=None, font=None, alignment=None, merge=None):
    cell = ws.cell(row=row, column=col, value=value)
    if fill:
        cell.fill = fill
    if font:
        cell.font = font
    if alignment:
        cell.alignment = alignment
    apply_border(cell)

    if merge:
        ws.merge_cells(start_row=row, start_column=col, end_row=merge[0], end_column=merge[1])

    return cell


def get_company_name(ws_in):
    company_name = "EMPRESA"
    for row in range(1, min(ws_in.max_row, 15) + 1):
        for col in range(1, min(ws_in.max_column, 8) + 1):
            val = ws_in.cell(row=row, column=col).value
            if val and "EMPRESA:" in normalize_text(val):
                company_name = str(val).replace("Empresa:", "").replace("EMPRESA:", "").strip()
                return company_name
    return company_name


def get_selected_employees(employee_index, employee_names):
    selected = []
    seen_rows = set()
    not_found = []

    for emp_name in employee_names:
        emp_name = emp_name.strip()
        if not emp_name:
            continue

        data = find_employee(employee_index, emp_name)
        if not data:
            not_found.append(emp_name)
            continue

        row_id = data["row"]
        if row_id not in seen_rows:
            selected.append(data)
            seen_rows.add(row_id)

    return selected, not_found


def group_days_into_weeks(year, month):
    first_day = date(year, month, 1)

    if first_day.weekday() <= 5:
        start_monday = first_day - timedelta(days=first_day.weekday())
    else:
        start_monday = first_day + timedelta(days=1)

    weeks = []
    current_monday = start_monday

    for _ in range(4):
        week = []
        for i in range(6):  # Lunes a sábado
            week.append(current_monday + timedelta(days=i))
        weeks.append(week)
        current_monday += timedelta(days=7)

    return weeks


def set_range_border(ws, start_row, start_col, end_row, end_col):
    for r in range(start_row, end_row + 1):
        for c in range(start_col, end_col + 1):
            apply_border(ws.cell(r, c))

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1 cm = 567 twips aprox
    tcW.set(qn('w:type'), 'dxa')


def set_row_height(row, height_twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def center_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), 'center')

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color_hex)


def set_cell_text(cell, text, bold=False, size=9, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def style_word_cell(
    cell,
    text="",
    bg_color="FFFFFF",
    bold=False,
    size=9,
    align="center",
    font_color="000000"
):
    cell.text = ""

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(font_color)

    r = run._element
    rPr = r.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_cell_background(cell, bg_color)

    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": "000000"},
        bottom={"val": "single", "sz": "8", "color": "000000"},
        left={"val": "single", "sz": "8", "color": "000000"},
        right={"val": "single", "sz": "8", "color": "000000"},
    )

import colorsys


def hsla_to_hex(h, s, l, a=1):
    """
    Convierte HSLA a HEX RGB.
    h: 0-360
    s: 0-100
    l: 0-100
    a: se acepta por compatibilidad, pero no se usa porque Word/Excel no manejan alpha en rellenos simples.
    """
    h = max(0, min(360, h)) / 360
    s = max(0, min(100, s)) / 100
    l = max(0, min(100, l)) / 100

    r, g, b = colorsys.hls_to_rgb(h, l, s)
    return "{:02X}{:02X}{:02X}".format(int(r * 255), int(g * 255), int(b * 255))


def get_theme_palettes():
    """
    Paletas definidas en HSL/HSLA.
    Puedes agregar más después.
    """
    return {
        "Rojo clásico": {
            "primary": (0, 100, 50, 1),        # rojo fuerte
            "secondary": (22, 37, 83, 1),      # beige rosado
            "border": (0, 0, 0, 1),            # negro
            "title": (0, 100, 50, 1),          # rojo título
            "text_light": (0, 0, 100, 1),      # blanco
            "text_dark": (0, 0, 0, 1),         # negro
            "neutral": (0, 0, 100, 1),         # blanco
        },
        "Azul corporativo": {
            "primary": (210, 100, 36, 1),
            "secondary": (210, 35, 86, 1),
            "border": (210, 25, 22, 1),
            "title": (210, 100, 36, 1),
            "text_light": (0, 0, 100, 1),
            "text_dark": (0, 0, 0, 1),
            "neutral": (0, 0, 100, 1),
        },
        "Verde institucional": {
            "primary": (145, 65, 34, 1),
            "secondary": (145, 30, 85, 1),
            "border": (145, 20, 20, 1),
            "title": (145, 65, 34, 1),
            "text_light": (0, 0, 100, 1),
            "text_dark": (0, 0, 0, 1),
            "neutral": (0, 0, 100, 1),
        },
        "Dorado sobrio": {
            "primary": (43, 100, 47, 1),
            "secondary": (43, 55, 88, 1),
            "border": (35, 40, 28, 1),
            "title": (43, 100, 37, 1),
            "text_light": (0, 0, 100, 1),
            "text_dark": (0, 0, 0, 1),
            "neutral": (0, 0, 100, 1),
        },
    }


def build_theme_from_palette_name(palette_name):
    palettes = get_theme_palettes()

    # color por default si el usuario no selecciona nada
    if not palette_name or palette_name not in palettes:
        palette_name = "Rojo clásico"

    raw = palettes[palette_name]

    return {
        "name": palette_name,
        "primary": hsla_to_hex(*raw["primary"]),
        "secondary": hsla_to_hex(*raw["secondary"]),
        "border": hsla_to_hex(*raw["border"]),
        "title": hsla_to_hex(*raw["title"]),
        "text_light": hsla_to_hex(*raw["text_light"]),
        "text_dark": hsla_to_hex(*raw["text_dark"]),
        "neutral": hsla_to_hex(*raw["neutral"]),
    }


def init_theme_session():
    if "selected_palette_name" not in st.session_state:
        st.session_state["selected_palette_name"] = "Rojo clásico"

    if "theme" not in st.session_state:
        st.session_state["theme"] = build_theme_from_palette_name(
            st.session_state["selected_palette_name"]
        )


def render_global_theme_selector():
    init_theme_session()

    palettes = list(get_theme_palettes().keys())

    st.sidebar.markdown("## Diseño del documento")

    selected = st.sidebar.selectbox(
        "Paleta global",
        palettes,
        index=palettes.index(st.session_state["selected_palette_name"]),
        key="global_palette_select"
    )

    st.session_state["selected_palette_name"] = selected
    st.session_state["theme"] = build_theme_from_palette_name(selected)

    theme = st.session_state["theme"]

    st.sidebar.markdown(
        f"""
        <div style="display:flex; gap:8px; margin-top:8px; margin-bottom:8px;">
            <div style="width:28px; height:28px; background:#{theme['primary']}; border:1px solid #000;"></div>
            <div style="width:28px; height:28px; background:#{theme['secondary']}; border:1px solid #000;"></div>
            <div style="width:28px; height:28px; background:#{theme['border']}; border:1px solid #000;"></div>
            <div style="width:28px; height:28px; background:#{theme['neutral']}; border:1px solid #000;"></div>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.caption(f"Paleta activa: {theme['name']}")


def get_active_theme():
    init_theme_session()
    return st.session_state["theme"]

def configure_document_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

def add_month_table_to_doc(doc, company_name, year, month, selected_employees, theme):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BITÁCORA DE ASISTENCIA")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string(theme["title"])

    doc.add_paragraph("")

    month_name = f"{MONTH_NAMES_ES[month]} {year}"
    weeks = group_days_into_weeks(year, month)
    official_holidays = get_official_mexico_holidays(year)
    first_day_of_month = date(year, month, 1)

    # 4 filas superiores de títulos
    # 2 filas de encabezados
    # n filas de empleados
    total_cols = 26
    total_rows = 5 + len(selected_employees)

    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    center_table(table)

    col_widths = [2.1, 5.2] + [0.72] * 24

    for row in table.rows:
        for i, w in enumerate(col_widths):
            set_cell_width(row.cells[i], w)

    # Alturas de filas
    set_row_height(table.rows[0], 520)  # empresa
    set_row_height(table.rows[1], 520)  # mes
    set_row_height(table.rows[2], 520)  # lista
    set_row_height(table.rows[3], 500)  # semanas
    set_row_height(table.rows[4], 650)  # dias
    # ===== TÍTULOS DENTRO DE LA TABLA =====

    company_cell = table.cell(0, 0)
    company_cell.merge(table.cell(0, total_cols - 1))
    style_word_cell(
        company_cell,
        company_name,
        bg_color=theme["primary"],
        bold=True,
        size=13,
        font_color=theme["text_light"]
    )

    month_cell = table.cell(1, 0)
    month_cell.merge(table.cell(1, total_cols - 1))
    style_word_cell(
        month_cell,
        month_name,
        bg_color=theme["neutral"],
        bold=True,
        size=12,
        font_color=theme["text_dark"]
    )

    list_cell = table.cell(2, 0)
    list_cell.merge(table.cell(2, total_cols - 1))
    style_word_cell(
        list_cell,
        "LISTA DE ASISTENCIA",
        bg_color=theme["primary"],
        bold=True,
        size=12,
        font_color=theme["text_light"]
    )

    # ===== ENCABEZADOS =====
    header_row_1 = 3
    header_row_2 = 4

    cell_num = table.cell(header_row_1, 0)
    cell_num.merge(table.cell(header_row_2, 0))
    style_word_cell(cell_num, "NUM DE\nEMPLEADO", bg_color="FFFFFF", bold=True, size=9)

    cell_name = table.cell(header_row_1, 1)
    cell_name.merge(table.cell(header_row_2, 1))
    style_word_cell(cell_name, "NOMBRE\nCOMPLETO", bg_color="FFFFFF", bold=True, size=9)

    current_col = 2
    for i, week in enumerate(weeks[:4], start=1):
        start_col = current_col
        end_col = current_col + 5

        week_cell = table.cell(header_row_1, start_col)
        week_cell.merge(table.cell(header_row_1, end_col))
        style_word_cell(
            week_cell,
            f"SEMANA {i}",
            bg_color=theme["secondary"],
            bold=True,
            size=10,
            font_color=theme["text_dark"]
        )

        for offset, day in enumerate(week):
            style_word_cell(
                table.cell(header_row_2, current_col + offset),
                DAY_LETTERS[day.weekday()],
                bg_color="FFFFFF",
                bold=True,
                size=9
            )

        current_col += 6

    # ===== FILAS DE EMPLEADOS =====
    data_start_row = 5

    for idx, emp in enumerate(selected_employees, start=1):
        row_idx = data_start_row + (idx - 1)
        set_row_height(table.rows[row_idx], 620)

        employee_name = emp.get("clean_name", emp.get("original_name", ""))
        pre_month_status = get_pre_month_default_status(emp["statuses"], year, month)

        style_word_cell(table.cell(row_idx, 0), str(idx), bg_color="FFFFFF", bold=True, size=11)
        style_word_cell(table.cell(row_idx, 1), employee_name, bg_color="FFFFFF", bold=False, size=9, align="left")

        current_col = 2
        for week in weeks[:4]:
            for day in week:
                if day < first_day_of_month:
                    fill_color = STATUS_COLORS.get(normalize_text(pre_month_status), "FFFFFF")
                elif day.year == year and day.month == month:
                    status = normalize_text(emp["statuses"].get(day, ""))

                    if day in official_holidays:
                        fill_color = STATUS_COLORS.get("INHABIL", "FF0000")
                    elif status == "NO HABIA INGRESADO":
                        fill_color = STATUS_COLORS.get("NO HABÍA INGRESADO", "FFFFFF")
                    else:
                        fill_color = STATUS_COLORS.get(status, "FFFFFF")
                else:
                    fill_color = "FFFFFF"

                style_word_cell(table.cell(row_idx, current_col), "", bg_color=fill_color, size=8)
                current_col += 1

    doc.add_paragraph("")

    # Leyenda en una sola fila
    legend = doc.add_table(rows=1, cols=len(STATUS_LABELS))
    legend.style = "Table Grid"
    legend.autofit = False
    legend.alignment = WD_TABLE_ALIGNMENT.CENTER
    center_table(legend)

    legend_widths = [2.3] * len(STATUS_LABELS)
    for row in legend.rows:
        for i, w in enumerate(legend_widths):
            set_cell_width(row.cells[i], w)

    set_row_height(legend.rows[0], 520)

    for i, (label, color) in enumerate(STATUS_LABELS):
        style_word_cell(
            legend.cell(0, i),
            label,
            bg_color=color,
            bold=False,
            size=8
        )

    doc.add_paragraph("")

def generate_report_word(input_file, employee_names, output_file, theme):
    wb_in = load_workbook(input_file, data_only=True)
    ws_in = wb_in.active

    header_row, empleado_col, cliente_col, date_cols = find_header_row_and_columns(ws_in)
    employee_index = build_employee_index(ws_in, header_row, empleado_col, cliente_col, date_cols)

    company_name = get_company_name(ws_in)
    selected_employees, not_found = get_selected_employees(employee_index, employee_names)

    if not selected_employees:
        raise ValueError("No se encontró ningún empleado con los nombres capturados.")

    all_dates = sorted(selected_employees[0]["statuses"].keys())
    months_present = get_months_present(all_dates)

    if not months_present:
        raise ValueError("No se encontraron fechas válidas en el archivo.")

    bimester_months = months_present[:2]

    doc = Document()
    configure_document_landscape(doc)

    for i, (year, month) in enumerate(bimester_months):
        if i > 0:
            doc.add_page_break()
        add_month_table_to_doc(doc, company_name, year, month, selected_employees, theme)

    doc.save(output_file)
    return not_found


def create_month_table(ws_out, start_row, company_name, year, month, selected_employees):
    red_fill = PatternFill(fill_type="solid", fgColor="FF0000")
    light_header_fill = PatternFill(fill_type="solid", fgColor="EAD8C8")
    white_fill = PatternFill(fill_type="solid", fgColor="FFFFFF")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    bold_white = Font(bold=True, color="FFFFFF", size=12)
    bold_black = Font(bold=True, color="000000", size=11)
    title_font = Font(bold=True, color="FF0000", size=18)
    normal_font = Font(size=11, color="000000")

    widths = {
        1: 4,
        2: 16,
        3: 28,
    }
    for c in range(4, 40):
        widths[c] = 5

    for col, width in widths.items():
        ws_out.column_dimensions[get_column_letter(col)].width = width

    weeks = group_days_into_weeks(year, month)

    start_col = 4
    col = start_col
    week_ranges = []

    for week in weeks[:4]:
        end_col = col + 6 - 1
        week_ranges.append((col, end_col, week))
        col = end_col + 1

    last_col = max(14, week_ranges[-1][1])

    if start_row == 2:
        ws_out.merge_cells(start_row=start_row, start_column=2, end_row=start_row, end_column=last_col)
        ws_out.cell(start_row, 2, "BITÁCORA DE ASISTENCIA")
        ws_out.cell(start_row, 2).font = title_font
        ws_out.cell(start_row, 2).alignment = center

    ws_out.merge_cells(start_row=start_row + 2, start_column=2, end_row=start_row + 2, end_column=last_col)
    set_cell(ws_out, start_row + 2, 2, company_name, fill=red_fill, font=bold_white, alignment=center)

    month_name = f"{MONTH_NAMES_ES[month]} {year}"
    ws_out.merge_cells(start_row=start_row + 3, start_column=2, end_row=start_row + 3, end_column=last_col)
    set_cell(ws_out, start_row + 3, 2, month_name, fill=white_fill, font=Font(bold=True, size=14), alignment=center)

    ws_out.merge_cells(start_row=start_row + 4, start_column=2, end_row=start_row + 4, end_column=last_col)
    set_cell(ws_out, start_row + 4, 2, "LISTA DE ASISTENCIA", fill=red_fill, font=bold_white, alignment=center)

    header_row_1 = start_row + 5
    header_row_2 = start_row + 6

    ws_out.merge_cells(start_row=header_row_1, start_column=2, end_row=header_row_2, end_column=2)
    ws_out.merge_cells(start_row=header_row_1, start_column=3, end_row=header_row_2, end_column=3)

    set_cell(ws_out, header_row_1, 2, "NUM DE\nEMPLEADO", fill=white_fill, font=Font(bold=True, size=11), alignment=center)
    set_cell(ws_out, header_row_1, 3, "NOMBRE\nCOMPLETO", fill=white_fill, font=Font(bold=True, size=11), alignment=center)

    for i, (week_start, week_end, week_days) in enumerate(week_ranges, start=1):
        ws_out.merge_cells(start_row=header_row_1, start_column=week_start, end_row=header_row_1, end_column=week_end)
        set_cell(ws_out, header_row_1, week_start, f"SEMANA {i}", fill=light_header_fill, font=bold_black, alignment=center)
        set_range_border(ws_out, header_row_1, week_start, header_row_1, week_end)

        for offset, day in enumerate(week_days):
            cell = ws_out.cell(row=header_row_2, column=week_start + offset, value=DAY_LETTERS[day.weekday()])
            cell.font = bold_black
            cell.alignment = center
            apply_border(cell)

    data_start_row = start_row + 7
    current_row = data_start_row

    for index, emp in enumerate(selected_employees, start=1):
        employee_number = str(index)
        employee_name = emp.get("clean_name", emp.get("original_name", ""))

        set_cell(ws_out, current_row, 2, employee_number, fill=white_fill, font=Font(bold=True), alignment=center)
        set_cell(ws_out, current_row, 3, employee_name, fill=white_fill, font=normal_font, alignment=left_align)

        first_day_of_month = date(year, month, 1)
        official_holidays = get_official_mexico_holidays(year)
        pre_month_status = get_pre_month_default_status(emp["statuses"], year, month)

        for week_start, week_end, week_days in week_ranges:
            for offset, day in enumerate(week_days):
                if day < first_day_of_month:
                    fill = status_to_fill(pre_month_status)

                elif day.year == year and day.month == month:
                    status = normalize_text(emp["statuses"].get(day, ""))

                    if day in official_holidays:
                        fill = status_to_fill("INHÁBIL")
                    elif status == "NO HABIA INGRESADO":
                        fill = status_to_fill("NO HABÍA INGRESADO")
                    else:
                        fill = status_to_fill(status)

                else:
                    fill = white_fill

                set_cell(ws_out, current_row, week_start + offset, "", fill=fill, alignment=center)

        current_row += 1

    legend_row = current_row + 2
    legend_col = 2

    for label, color in STATUS_LABELS:
        ws_out.merge_cells(
            start_row=legend_row,
            start_column=legend_col,
            end_row=legend_row + 1,
            end_column=legend_col + 1
        )
        c = ws_out.cell(row=legend_row, column=legend_col, value=label)
        c.fill = PatternFill(fill_type="solid", fgColor=color)
        c.alignment = center
        c.font = Font(bold=False, color="000000")
        apply_border(c)

        for r in range(legend_row, legend_row + 2):
            for cc in range(legend_col, legend_col + 2):
                apply_border(ws_out.cell(r, cc))

        legend_col += 3

    ws_out.row_dimensions[start_row].height = 28
    ws_out.row_dimensions[start_row + 2].height = 24
    ws_out.row_dimensions[start_row + 3].height = 24
    ws_out.row_dimensions[start_row + 4].height = 24
    ws_out.row_dimensions[header_row_1].height = 42
    ws_out.row_dimensions[header_row_2].height = 24

    for r in range(data_start_row, current_row):
        ws_out.row_dimensions[r].height = 38

    return legend_row + 4


def generate_report(input_file, employee_names, output_file):
    wb_in = load_workbook(input_file, data_only=True)
    ws_in = wb_in.active

    header_row, empleado_col, cliente_col, date_cols = find_header_row_and_columns(ws_in)
    employee_index = build_employee_index(ws_in, header_row, empleado_col, cliente_col, date_cols)

    company_name = get_company_name(ws_in)
    selected_employees, not_found = get_selected_employees(employee_index, employee_names)

    if not selected_employees:
        raise ValueError("No se encontró ningún empleado con los nombres capturados.")

    all_dates = sorted(selected_employees[0]["statuses"].keys())
    months_present = get_months_present(all_dates)

    if not months_present:
        raise ValueError("No se encontraron fechas válidas en el archivo.")

    bimester_months = months_present[:2]

    wb_out = Workbook()
    ws_out = wb_out.active
    ws_out.title = "Bitácora"

    next_row = 2
    first_table = True

    for year, month in bimester_months:
        if not first_table:
            next_row += 2

        next_row = create_month_table(
            ws_out=ws_out,
            start_row=next_row,
            company_name=company_name,
            year=year,
            month=month,
            selected_employees=selected_employees
        )

        first_table = False

    ws_out.freeze_panes = "D8"
    wb_out.save(output_file)
    return not_found


def mostrar_modulo_bitacora():
    st.title("Generador de Bitácora de Asistencia")

    st.write(
        "Escribe uno o varios nombres exactamente como aparecen en el Excel. "
        "También puede encontrar coincidencias parciales si solo coincide un empleado."
    )

    empleados_texto = st.text_area(
        "Nombres de empleados",
        height=180,
        placeholder="Empleado 1\nEmpleado 2\nEmpleado 3",
        key="bitacora_empleados"
    )

    archivo_excel = st.file_uploader(
        "Selecciona el archivo de asistencias",
        type=["xlsx", "xlsm", "xltx", "xltm"],
        key="bitacora_uploader"
    )

    nombre_salida = st.text_input(
        "Nombre del archivo de salida",
        value="bitacora_asistencia.docx",
        key="bitacora_salida"
    )

    theme = get_active_theme()

    if st.button("Generar Word", key="bitacora_generar"):
        if not archivo_excel:
            st.error("Selecciona el archivo de asistencias.")
            return

        employee_names = [x.strip() for x in empleados_texto.splitlines() if x.strip()]
        if not employee_names:
            st.error("Escribe al menos un nombre de empleado.")
            return

        if not nombre_salida.strip():
            nombre_salida = "bitacora_asistencia.docx"

        if not nombre_salida.lower().endswith(".docx"):
            nombre_salida += ".docx"

        input_path = None
        output_path = None

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_in:
                temp_in.write(archivo_excel.getbuffer())
                input_path = temp_in.name

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_out:
                output_path = temp_out.name

            not_found = generate_report_word(input_path, employee_names, output_path, theme)

            with open(output_path, "rb") as f:
                output_bytes = f.read()

            st.success("Archivo Word generado correctamente.")

            if not_found:
                st.warning(
                    "No se encontraron estos empleados:\n- " + "\n- ".join(not_found)
                )

            st.download_button(
                label="Descargar reporte Word",
                data=output_bytes,
                file_name=nombre_salida,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Ocurrió un problema:\n\n{str(e)}")

        finally:
            try:
                if input_path and os.path.exists(input_path):
                    os.remove(input_path)
            except Exception:
                pass

            try:
                if output_path and os.path.exists(output_path):
                    os.remove(output_path)
            except Exception:
                pass