from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def agregar_parrafo_con_negritas(doc, partes, fuente, size, color_rgb):
    p = doc.add_paragraph()

    for texto, bold in partes:
        run = p.add_run(texto)
        run.font.name = fuente
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(*color_rgb)

    return p


def poner_bordes_tabla(tabla):
    tbl = tabla._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")

    for nombre in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{nombre}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        border.set(qn("w:space"), "0")
        borders.append(border)

    tblPr.append(borders)


def colorear_celda(celda, color):
    tcPr = celda._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def configurar_texto_celda(
    celda,
    texto,
    fuente,
    size,
    bold,
    color_rgb,
    align=WD_ALIGN_PARAGRAPH.LEFT
):
    celda.text = ""

    p = celda.paragraphs[0]
    p.alignment = align

    run = p.add_run(str(texto))
    run.font.name = fuente
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)