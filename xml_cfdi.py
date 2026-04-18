import os
import tempfile
import xml.etree.ElementTree as ET
from io import BytesIO
from datetime import datetime

import pandas as pd
import streamlit as st
from openpyxl.styles import Font

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def limpiar_concepto(texto):
    if not texto:
        return ""
    texto = texto.strip()
    if texto.startswith("-"):
        texto = texto.lstrip("-").strip()
    return texto


def formatear_fecha(fecha_raw):
    if not fecha_raw:
        return ""
    try:
        fecha_dt = datetime.fromisoformat(fecha_raw)
        return f"{fecha_dt.day:02d}/{fecha_dt.month:02d}/{fecha_dt.year}"
    except Exception:
        try:
            solo_fecha = fecha_raw.split("T")[0]
            anio, mes, dia = solo_fecha.split("-")
            return f"{int(dia):02d}/{int(mes):02d}/{anio}"
        except Exception:
            return fecha_raw


def parse_xml(file):
    registros = []

    try:
        tree = ET.parse(file)
        root = tree.getroot()

        ns = {
            "cfdi": "http://www.sat.gob.mx/cfd/4",
            "tfd": "http://www.sat.gob.mx/TimbreFiscalDigital"
        }

        fecha_raw = root.attrib.get("Fecha", "")
        total = root.attrib.get("Total", "0")
        serie = root.attrib.get("Serie", "")
        folio = root.attrib.get("Folio", "")
        serie = str(serie).strip()
        folio = str(folio).strip()
        serie_folio = f"{serie}-{folio}" if serie and folio else (serie or folio)

        fecha = formatear_fecha(fecha_raw)

        try:
            importe_total = float(total)
        except Exception:
            importe_total = 0.0

        timbre = root.find(".//tfd:TimbreFiscalDigital", ns)
        uuid = ""
        if timbre is not None:
            uuid = timbre.attrib.get("UUID", "")

        conceptos = root.findall(".//cfdi:Concepto", ns)

        for concepto in conceptos:
            descripcion = limpiar_concepto(concepto.attrib.get("Descripcion", ""))

            registros.append({
                "CONCEPTO": descripcion,
                "FECHA": fecha,
                "SERIE Y FOLIO": serie_folio,
                "UUID": uuid,
                "IMPORTE TOTAL": importe_total
            })

    except Exception as e:
        st.error(f"Error al leer el archivo {getattr(file, 'name', 'XML')}: {e}")

    return registros


def generar_excel_xml(df):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Datos")

        worksheet = writer.sheets["Datos"]

        for cell in worksheet[1]:
            cell.font = Font(bold=True)

        for col in worksheet.iter_cols(min_col=5, max_col=5, min_row=2):
            for cell in col:
                cell.number_format = '$#,##0.00'

        worksheet.column_dimensions["A"].width = 90
        worksheet.column_dimensions["B"].width = 15
        worksheet.column_dimensions["C"].width = 22
        worksheet.column_dimensions["D"].width = 40
        worksheet.column_dimensions["E"].width = 18

    output.seek(0)
    return output


# =========================
# Helpers Word
# =========================

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_row_height(row, height_twips):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def center_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")


def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color_hex)


def set_cell_border(cell, top="F2B000", bottom="F2B000", left="F2B000", right="F2B000", size="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge, color in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        tag = f"w:{edge}"
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)

def set_cell_no_wrap(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    no_wrap = tcPr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tcPr.append(no_wrap)


def get_content_width_cm(doc):
    section = doc.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 360000

def style_word_cell(
    cell,
    text="",
    bg_color="FFFFFF",
    font_color="000000",
    bold=False,
    size=10,
    align="center",
    border_color="F2B000",
    no_wrap=False
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(font_color)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_background(cell, bg_color)
    set_cell_border(
        cell,
        top=border_color,
        bottom=border_color,
        left=border_color,
        right=border_color
    )

    if no_wrap:
        set_cell_no_wrap(cell)


def configure_document_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)


def obtener_rango_meses(df):
    fechas_validas = []
    for valor in df["FECHA"].tolist():
        try:
            fecha = datetime.strptime(str(valor), "%d/%m/%Y")
            fechas_validas.append(fecha)
        except Exception:
            pass

    if not fechas_validas:
        return "CONTROL DE PAGOS"

    fechas_validas.sort()
    meses_es = {
        1: "ENERO", 2: "FEBRERO", 3: "MARZO", 4: "ABRIL",
        5: "MAYO", 6: "JUNIO", 7: "JULIO", 8: "AGOSTO",
        9: "SEPTIEMBRE", 10: "OCTUBRE", 11: "NOVIEMBRE", 12: "DICIEMBRE"
    }

    f1 = fechas_validas[0]
    f2 = fechas_validas[-1]

    if f1.month == f2.month and f1.year == f2.year:
        return f"{meses_es[f1.month]} {f1.year}"

    if f1.year == f2.year:
        return f"{meses_es[f1.month]}-{meses_es[f2.month]} {f1.year}"

    return f"{meses_es[f1.month]} {f1.year} - {meses_es[f2.month]} {f2.year}"


def generar_word_xml(df, output_file):
    doc = Document()
    configure_document_landscape(doc)

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTROL DE PAGOS")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("B8860B")

    rango = obtener_rango_meses(df)

    # Tabla con 5 columnas reales
    headers = ["CONCEPTO", "FECHA", "SERIE Y FOLIO", "UUID", "IMPORTE TOTAL"]
    total_rows = len(df) + 1
    total_cols = 5

    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    center_table(table)

    # Anchos más equilibrados para que no se partan títulos ni fechas
    col_widths = [7.9, 2.9, 3.6, 6.1, 3.6]

    for row in table.rows:
        for i, w in enumerate(col_widths):
            set_cell_width(row.cells[i], w)

    # Subtítulo alineado al inicio real de la tabla
    total_table_width = sum(col_widths)
    content_width = get_content_width_cm(doc)
    left_indent_cm = max((content_width - total_table_width) / 2, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(left_indent_cm)

    run = p.add_run(rango)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("F0B000")

    # Mover subtítulo arriba de la tabla
    table._element.addprevious(p._element)
    spacer = doc.add_paragraph("")
    table._element.addprevious(spacer._element)

    # Header
    header_bg = "F2B000"
    border_color = "F2B000"

    for i, h in enumerate(headers):
        style_word_cell(
            table.cell(0, i),
            h,
            bg_color=header_bg,
            font_color="000000",
            bold=True,
            size=11,
            border_color=border_color,
            no_wrap=True
        )

    set_row_height(table.rows[0], 430)

    # Data
    for idx, (_, row_data) in enumerate(df.iterrows(), start=1):
        set_row_height(table.rows[idx], 560)

        concepto = str(row_data["CONCEPTO"]) if pd.notna(row_data["CONCEPTO"]) else ""
        fecha = str(row_data["FECHA"]) if pd.notna(row_data["FECHA"]) else ""
        serie_folio = str(row_data["SERIE Y FOLIO"]) if pd.notna(row_data["SERIE Y FOLIO"]) else ""
        uuid = str(row_data["UUID"]) if pd.notna(row_data["UUID"]) else ""

        try:
            importe = f"${float(row_data['IMPORTE TOTAL']):,.2f}"
        except Exception:
            importe = str(row_data["IMPORTE TOTAL"])

        style_word_cell(
            table.cell(idx, 0),
            concepto,
            bg_color="FFFFFF",
            font_color="000000",
            bold=False,
            size=11,
            align="left",
            border_color=border_color
        )
        style_word_cell(
            table.cell(idx, 1),
            fecha,
            bg_color="FFFFFF",
            size=10,
            border_color=border_color,
            no_wrap=True
        )
        style_word_cell(
            table.cell(idx, 2),
            serie_folio,
            bg_color="FFFFFF",
            size=11,
            border_color=border_color,
            no_wrap=True
        )
        style_word_cell(
            table.cell(idx, 3),
            uuid,
            bg_color="FFFFFF",
            size=9,
            align="left",
            border_color=border_color
        )
        style_word_cell(
            table.cell(idx, 4),
            importe,
            bg_color="FFFFFF",
            size=11,
            border_color=border_color,
            no_wrap=True
        )

    doc.save(output_file)


def mostrar_modulo_xml():
    st.title("Lector de XML CFDI")
    st.write("Sube uno o varios archivos XML para extraer CONCEPTO, FECHA, SERIE Y FOLIO, UUID e IMPORTE TOTAL.")

    files = st.file_uploader(
        "Selecciona uno o varios archivos XML",
        type=["xml"],
        accept_multiple_files=True,
        key="xml_uploader"
    )

    if files:
        all_data = []

        for file in files:
            all_data.extend(parse_xml(file))

        if all_data:
            df = pd.DataFrame(all_data)

            df["_FECHA_ORDEN"] = pd.to_datetime(df["FECHA"], format="%d/%m/%Y", errors="coerce")
            df = df.sort_values(by="_FECHA_ORDEN", ascending=True, na_position="last").drop(columns=["_FECHA_ORDEN"])
            df = df.reset_index(drop=True)
            st.subheader("Vista previa")
            df_vista = df.copy()
            df_vista["IMPORTE TOTAL"] = df_vista["IMPORTE TOTAL"].apply(lambda x: f"${x:,.2f}")
            st.dataframe(df_vista, use_container_width=True)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
                tmp_doc_path = tmp_doc.name

            try:
                with st.spinner("Generando Word..."):
                    generar_word_xml(df, tmp_doc_path)

                with open(tmp_doc_path, "rb") as f:
                    word_bytes = f.read()

                st.download_button(
                    label="Descargar Word",
                    data=word_bytes,
                    file_name="resultado_xml.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            finally:
                if os.path.exists(tmp_doc_path):
                    os.remove(tmp_doc_path)

        else:
            st.warning("No se encontraron datos válidos en los archivos XML.")
    else:
        st.info("Sube uno o varios archivos XML para comenzar.")