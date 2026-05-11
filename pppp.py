import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from copy import deepcopy
from io import BytesIO
from datetime import date
import matplotlib.font_manager as fm
import os

#st.set_page_config(page_title="Generador de Resumen Ejecutivo", layout="centered")

RUTA_CONTENIDO = os.path.join("datosp", "6_RESUMEN_EJECUTIVO_AJUSTADO_FUNCIONAL.docx")

MESES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre"
}

def fecha_larga_es(fecha):
    return f"{fecha.day} de {MESES_ES[fecha.month]} del {fecha.year}"


def mostrar_modulo_resumen():
    st.title("Generador de Resumen Ejecutivo")

    plantilla_subida = st.file_uploader(
        "Sube la plantilla Word con encabezado/diseño",
        type=["docx"]
    )

    empresa_recibe = st.text_input("Empresa que recibe el servicio")
    empresa_brinda = st.text_input("Empresa que brinda el servicio")
    nombre_programa = st.text_area("Nombre del programa")

    conceptos = st.text_area(
        "Conceptos desarrollados",
        value="Plan integral que permita efectuar con nuevas tecnologías el monitoreo de rutas logísticas.\nDigitalización de trámites para reducir tiempos y costos en los procesos de distribución.",
        help="Escribe un concepto por línea."
    )

    st.subheader("Periodo del servicio")

    col_fecha1, col_fecha2 = st.columns(2)

    with col_fecha1:
        fecha_inicio = st.date_input(
            "Fecha de inicio del servicio",
            value=date(2025, 9, 29),
            format="DD/MM/YYYY"
        )

    with col_fecha2:
        fecha_termino = st.date_input(
            "Fecha de término del servicio",
            value=date(2026, 3, 17),
            format="DD/MM/YYYY"
        )

    fecha_servicio = f"{fecha_larga_es(fecha_inicio)} al {fecha_larga_es(fecha_termino)}"

    fuentes = sorted(set([f.name for f in fm.fontManager.ttflist]))

    st.subheader("Configuración de estilos del documento")

    col1, col2 = st.columns(2)

    with col1:
        fuente_titulo1 = st.selectbox(
            "Fuente Título 1",
            fuentes,
            index=fuentes.index("Arial") if "Arial" in fuentes else 0,
            key="fuente_titulo1"
        )
        tamano_titulo1 = st.number_input("Tamaño Título 1", 6, 80, 22)
        color_titulo1 = st.color_picker("Color Título 1", "#58005A")

        fuente_titulo2 = st.selectbox(
            "Fuente Título 2",
            fuentes,
            index=fuentes.index("Arial") if "Arial" in fuentes else 0,
            key="fuente_titulo2"
        )
        tamano_titulo2 = st.number_input("Tamaño Título 2", 6, 80, 18)
        color_titulo2 = st.color_picker("Color Título 2", "#000000")

    with col2:
        fuente_titulo3 = st.selectbox(
            "Fuente Título 3",
            fuentes,
            index=fuentes.index("Arial") if "Arial" in fuentes else 0,
            key="fuente_titulo3"
        )
        tamano_titulo3 = st.number_input("Tamaño Título 3", 6, 80, 14)
        color_titulo3 = st.color_picker("Color Título 3", "#000000")

        fuente_normal = st.selectbox(
            "Fuente Texto normal",
            fuentes,
            index=fuentes.index("Arial") if "Arial" in fuentes else 0,
            key="fuente_normal"
        )
        tamano_normal = st.number_input("Tamaño Texto normal", 6, 60, 11)
        color_normal = st.color_picker("Color Texto normal", "#000000")


    def hex_a_rgb(hex_color):
        hex_color = hex_color.replace("#", "")
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )


    def aplicar_fuente(run, fuente, tamano, color_hex, negrita=False):
        run.font.name = fuente
        run.font.size = Pt(tamano)
        run.font.color.rgb = hex_a_rgb(color_hex)
        run.bold = negrita

        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.rFonts

        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)

        rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", fuente)
        rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", fuente)
        rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", fuente)
        rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs", fuente)


    def obtener_formato_por_estilo(parrafo, configuracion):
        texto = parrafo.text.strip().upper()

        try:
            nombre_estilo = parrafo.style.name
        except:
            nombre_estilo = ""

        if texto == "RESUMEN EJECUTIVO":
            return configuracion["titulo1"]

        if texto == "{{EMPRESA_BRINDA}}":
            return configuracion["titulo2"]

        if texto == "{{NOMBRE_PROGRAMA}}":
            return configuracion["titulo3"]

        if nombre_estilo in ["Título 1", "Heading 1"]:
            return configuracion["titulo1"]

        elif nombre_estilo in ["Título 2", "Heading 2"]:
            return configuracion["titulo2"]

        elif nombre_estilo in ["Título 3", "Heading 3"]:
            return configuracion["titulo3"]

        else:
            return configuracion["normal"]


    def limpiar_runs(parrafo):
        for run in parrafo.runs:
            run.text = ""


    def escribir_texto_en_parrafo(parrafo, texto, fuente, tamano, color, negrita=False):
        limpiar_runs(parrafo)
        run = parrafo.add_run(texto)
        aplicar_fuente(run, fuente, tamano, color, negrita)


    def reemplazar_en_parrafo(parrafo, reemplazos, configuracion):
        if not parrafo.runs:
            return

        fuente, tamano, color = obtener_formato_por_estilo(parrafo, configuracion)

        for run in parrafo.runs:
            texto = run.text

            if not texto:
                continue

            es_variable = False

            for clave, valor in reemplazos.items():
                if clave in texto:
                    texto = texto.replace(clave, valor)
                    es_variable = True

            run.text = texto

            aplicar_fuente(
                run,
                fuente,
                tamano,
                color,
                negrita=es_variable or bool(run.bold)
            )


    def reemplazar_conceptos_en_parrafo(parrafo, lista_conceptos, configuracion):
        texto = parrafo.text.strip()

        if not texto:
            return False

        fuente, tamano, color = configuracion["normal"]

        texto_conceptos = "\n".join(
            [f"{i + 1}.     {concepto}" for i, concepto in enumerate(lista_conceptos)]
        )

        if "{{CONCEPTOS_DESARROLLADOS}}" in texto:
            escribir_texto_en_parrafo(parrafo, texto_conceptos, fuente, tamano, color, False)
            return True

        if "Plan integral que permita efectuar" in texto:
            escribir_texto_en_parrafo(parrafo, texto_conceptos, fuente, tamano, color, False)
            return True

        if "Digitalización de trámites" in texto:
            escribir_texto_en_parrafo(parrafo, "", fuente, tamano, color, False)
            return True

        return False


    def procesar_documento(doc, reemplazos, configuracion, lista_conceptos):
        for parrafo in doc.paragraphs:
            reemplazar_conceptos_en_parrafo(parrafo, lista_conceptos, configuracion)
            reemplazar_en_parrafo(parrafo, reemplazos, configuracion)

        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        reemplazar_conceptos_en_parrafo(parrafo, lista_conceptos, configuracion)
                        reemplazar_en_parrafo(parrafo, reemplazos, configuracion)

        return doc


    def procesar_textboxes_xml(doc, reemplazos):
        for t in doc._element.xpath(".//*[local-name()='t']"):
            if t.text:
                for clave, valor in reemplazos.items():
                    t.text = t.text.replace(clave, valor)


    def limpiar_cuerpo_documento(doc):
        body = doc._body._element

        for elemento in list(body):
            if elemento.tag.endswith("sectPr"):
                continue

            body.remove(elemento)


    def insertar_contenido(destino, origen):
        body_destino = destino._body._element
        body_origen = origen._body._element

        sectPr = body_destino.sectPr

        for elemento in body_origen:
            if elemento.tag.endswith("sectPr"):
                continue

            nuevo_elemento = deepcopy(elemento)

            if sectPr is not None:
                body_destino.insert(body_destino.index(sectPr), nuevo_elemento)
            else:
                body_destino.append(nuevo_elemento)

        return destino


    if st.button("Generar Word", key="btn_generar_word_resumen"):

        if plantilla_subida is None:
            st.error("Primero sube la plantilla Word con encabezado/diseño.")

        elif not os.path.exists(RUTA_CONTENIDO):
            st.error(f"No se encontró el archivo: {RUTA_CONTENIDO}")

        elif not empresa_recibe or not empresa_brinda or not nombre_programa:
            st.error("Completa todos los campos.")

        elif not conceptos.strip():
            st.error("Agrega al menos un concepto.")

        elif fecha_inicio > fecha_termino:
            st.error("La fecha de inicio no puede ser posterior a la fecha de término.")

        else:
            plantilla = Document(plantilla_subida)
            contenido = Document(RUTA_CONTENIDO)

            lista_conceptos = [
                c.strip()
                for c in conceptos.splitlines()
                if c.strip()
            ]

            configuracion = {
                "titulo1": (fuente_titulo1, tamano_titulo1, color_titulo1),
                "titulo2": (fuente_titulo2, tamano_titulo2, color_titulo2),
                "titulo3": (fuente_titulo3, tamano_titulo3, color_titulo3),
                "normal": (fuente_normal, tamano_normal, color_normal)
            }

            reemplazos = {
                "{{EMPRESA_RECIBE}}": empresa_recibe.upper(),
                "{{EMPRESA_BRINDA}}": empresa_brinda.upper(),
                "{{NOMBRE_PROGRAMA}}": nombre_programa.upper(),
                "{{FECHA_SERVICIO}}": fecha_servicio,
                "29 de septiembre del 2025 al 17 de marzo del 2026": fecha_servicio
            }

            contenido = procesar_documento(
                contenido,
                reemplazos,
                configuracion,
                lista_conceptos
            )

            procesar_textboxes_xml(contenido, reemplazos)

            limpiar_cuerpo_documento(plantilla)

            documento_final = insertar_contenido(plantilla, contenido)

            salida = BytesIO()
            documento_final.save(salida)
            salida.seek(0)

            st.success("Documento generado correctamente.")

            st.download_button(
                label="Descargar Word generado",
                data=salida,
                file_name="resumen_ejecutivo_generado.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )